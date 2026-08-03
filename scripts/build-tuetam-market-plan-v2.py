from __future__ import annotations

from dataclasses import dataclass
from decimal import Decimal
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


WORKSPACE = Path(r"C:\Users\DELL\Desktop\Kyoto Masage")
OUT_DIR = WORKSPACE / "artifacts" / "tuetam-market-plan-v2.1"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = WORKSPACE / "artifacts" / "Tue_Tam_Care_Nghien_Cuu_Thi_Truong_Ke_Hoach_Kinh_Doanh_Tinh_Kha_Thi_v2.1.docx"
LOGO_PATH = WORKSPACE / "public" / "logo.png"
LOGO_SMALL_PATH = ASSET_DIR / "logo-512.png"

OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)


# Standard business brief preset, with a named Tuệ Tâm brand override for the
# cover, executive callouts and selected chart accents.
INK = "1F2937"
MUTED = "667085"
BLUE = "2E74B5"
NAVY = "1F4D78"
BRAND_RED = "B11927"
BRAND_GOLD = "B88922"
PALE_RED = "FDF2F2"
PALE_GOLD = "FCF7E8"
PALE_BLUE = "EFF6FD"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
WHITE = "FFFFFF"
GREEN = "18794E"
AMBER = "9A6700"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}

FX = Decimal("26100")
FIXED_SITE_TARGET = Decimal("2500000000")
BUSINESS_GMV_USD = Decimal("2000000")
BUSINESS_GMV_VND = BUSINESS_GMV_USD * FX
KTV_SHARE = Decimal("0.60")
TEAM_LEAD_SHARE = Decimal("0.05")
PLATFORM_SHARE = Decimal("0.20")
DISTRICT_DIRECTOR_SHARE = Decimal("0.05")
DIRECT_AFFILIATE_SHARE = Decimal("0.10")
ALLOCATION_TOTAL = KTV_SHARE + TEAM_LEAD_SHARE + PLATFORM_SHARE + DISTRICT_DIRECTOR_SHARE + DIRECT_AFFILIATE_SHARE
assert ALLOCATION_TOTAL == Decimal("1.00")


@dataclass(frozen=True)
class Source:
    code: str
    title: str
    publisher: str
    date: str
    url: str
    use: str


SOURCES = [
    Source(
        "S1",
        "Cấu hình vận hành Tuệ Tâm Care",
        "Tuệ Tâm Care",
        "26/07/2026",
        "",
        "Địa chỉ, số giường/ghế, số KTV, giờ mở cửa và các giả định mục tiêu do chủ dự án cung cấp.",
    ),
    Source(
        "S2",
        "Báo cáo kinh tế - xã hội Hà Nội quý I/2025",
        "Cổng thông tin điện tử Thành phố Hà Nội",
        "09/04/2025",
        "https://hanoi.gov.vn/kinh-te-xa-hoi-thu-do/bao-cao-tinh-hinh-kinh-te-xa-hoi-thang-ba-va-quy-i-nam-2025-cua-chi-cuc-thong-ke-ha-noi-425040909575038.htm",
        "Lực lượng lao động 4,154 triệu người; 4,114 triệu người có việc làm; 1,977 triệu lao động thành thị.",
    ),
    Source(
        "S3",
        "Tình hình đăng ký doanh nghiệp Hà Nội đến 31/05/2025",
        "Sở Tài chính Thành phố Hà Nội",
        "30/06/2025",
        "https://www.sotaichinh.hanoi.gov.vn/tinh-hinh-dang-ky-doanh-nghiep-tren-dia-ban-thanh-pho-ha-noi-thang-5-nam-2025-171610.html",
        "217.414 doanh nghiệp đang hoạt động; Hà Đông và Cầu Giấy thuộc nhóm có nhiều doanh nghiệp đăng ký mới.",
    ),
    Source(
        "S4",
        "TP.HCM đặt mục tiêu giải quyết việc làm cho 300.000 lao động năm 2025",
        "ITPC - Thành phố Hồ Chí Minh",
        "13/02/2025",
        "https://itpc.hochiminhcity.gov.vn/-/tp-hcm-at-muc-tieu-giai-quyet-viec-lam-cho-300-000-lao-ong-trong-nam-2025",
        "Dân số trung bình 9,54 triệu người và 4,73 triệu người có việc làm trong năm 2024.",
    ),
    Source(
        "S5",
        "Quản trị tài chính - nền tảng nâng cao năng lực cạnh tranh doanh nghiệp tư nhân",
        "ITPC - Thành phố Hồ Chí Minh",
        "2025",
        "https://itpc.hochiminhcity.gov.vn/-/quan-tri-tai-chinh-kiem-toan-thue-nen-tang-nang-cao-nang-luc-canh-tranh-doanh-nghiep-tu-nhan",
        "Khoảng 345.000 doanh nghiệp đang hoạt động tại TP.HCM, theo số liệu được bài viết dẫn từ Cục Thống kê TP.HCM.",
    ),
    Source(
        "S6",
        "Viet Nam Real Estate Market Report Q1/2025",
        "Savills Việt Nam",
        "2025",
        "https://vn.savills.com.vn/research_articles/164027/221117-0",
        "Hà Nội có 2,33 triệu m² nguồn cung văn phòng trong Q1/2025; phía Tây chiếm tỷ trọng lớn và nhu cầu hạng B duy trì tích cực.",
    ),
    Source(
        "S7",
        "The Global Wellness Economy: Country Rankings (2019-2022)",
        "Global Wellness Institute",
        "2024",
        "https://globalwellnessinstitute.org/wp-content/uploads/2024/01/2024GWI-Country-Rankings01212024.pdf",
        "Quy mô kinh tế wellness Việt Nam năm 2022 được ước tính 19,60 tỷ USD; chỉ dùng làm bối cảnh rộng, không dùng làm doanh thu dự báo.",
    ),
    Source(
        "S8",
        "Musculoskeletal health",
        "World Health Organization",
        "14/07/2022",
        "https://www.who.int/vietnam/news/fact-sheets/detail/musculoskeletal-conditions",
        "Các vấn đề cơ xương khớp ảnh hưởng khả năng làm việc và gây chi phí gián tiếp do vắng mặt hoặc giảm năng suất.",
    ),
    Source(
        "S9",
        "ILO and Viet Nam partners address psychosocial risks at work",
        "International Labour Organization",
        "12/05/2026",
        "https://www.ilo.org/resource/news/ilo-and-viet-nam-partners-address-psychosocial-risks-work-national",
        "Bối cảnh sức khỏe tâm lý - xã hội tại nơi làm việc và nhu cầu phòng ngừa rủi ro nghề nghiệp.",
    ),
    Source(
        "S10",
        "Tỷ giá ngoại tệ",
        "Vietcombank",
        "20/07/2026",
        "https://vietcombank.com.vn/vi-VN/To-chuc/Doanh-Nghi%E1%BB%87p/KHTC---Ty-gia",
        "Tỷ giá mua chuyển khoản USD/VND 26.110; mô hình dùng 26.100 VND/USD để hoạch định.",
    ),
    Source(
        "S11",
        "Nghị định 13/2023/NĐ-CP về bảo vệ dữ liệu cá nhân",
        "Cổng thông tin điện tử Chính phủ",
        "17/04/2023",
        "https://vanban.chinhphu.vn/default.aspx?docid=207759&pageid=27160",
        "Khung nghĩa vụ khi xử lý dữ liệu cá nhân, sự đồng ý, mục đích xử lý và bảo mật.",
    ),
]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def fmt_vnd(value: Decimal | float | int, decimals: int = 0) -> str:
    n = Decimal(str(value))
    if decimals == 0:
        return f"{int(n):,}".replace(",", ".") + " đ"
    return f"{float(n):,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".") + " đ"


def fmt_billion(value: Decimal | float | int, decimals: int = 2) -> str:
    n = Decimal(str(value)) / Decimal("1000000000")
    return f"{float(n):.{decimals}f}".replace(".", ",") + " tỷ đ"


def fmt_million(value: Decimal | float | int, decimals: int = 1) -> str:
    n = Decimal(str(value)) / Decimal("1000000")
    return f"{float(n):.{decimals}f}".replace(".", ",") + " triệu đ"


def safe_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_chart_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    title_font = safe_font(40, bold=True)
    sub_font = safe_font(23)
    draw.text((72, 48), title, fill="#1F2937", font=title_font)
    sub_lines = wrap_text(draw, subtitle, sub_font, 1450)
    y = 104
    for line in sub_lines:
        draw.text((74, y), line, fill="#667085", font=sub_font)
        y += 30


def save_fixed_site_chart(path: Path) -> None:
    scenarios = [
        ("Thận trọng", 8, 300_000, 1_728_000_000),
        ("Cơ sở", 10, 350_000, 2_520_000_000),
        ("Tăng trưởng", 14, 400_000, 4_032_000_000),
    ]
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw_chart_header(
        draw,
        "Kịch bản doanh thu hai cơ sở cố định",
        "Hai cơ sở, 360 ngày vận hành/năm; doanh thu chưa gồm tiền tip và chưa trừ chi phí.",
    )
    label_font = safe_font(26, bold=True)
    note_font = safe_font(20)
    value_font = safe_font(25, bold=True)
    colors = ["#7C8DA5", "#2E74B5", "#B11927"]
    max_value = max(x[3] for x in scenarios)
    x0, plot_w, y0 = 420, 1010, 235
    for idx, (label, visits, ticket, revenue) in enumerate(scenarios):
        y = y0 + idx * 185
        draw.text((70, y + 16), label, fill="#1F2937", font=label_font)
        detail = f"{visits} lượt/ngày/cơ sở × {ticket // 1000:,} nghìn đ".replace(",", ".")
        draw.text((70, y + 54), detail, fill="#667085", font=note_font)
        width = int(plot_w * revenue / max_value)
        draw.rounded_rectangle((x0, y, x0 + width, y + 92), radius=12, fill=colors[idx])
        value = f"{revenue / 1_000_000_000:.2f} tỷ đ".replace(".", ",")
        bbox = draw.textbbox((0, 0), value, font=value_font)
        tx = x0 + width - (bbox[2] - bbox[0]) - 22
        fill = "white" if width > 280 else "#1F2937"
        if width <= 280:
            tx = x0 + width + 15
        draw.text((tx, y + 28), value, fill=fill, font=value_font)
    draw.text((72, 818), "Nguồn: mô hình hoạch định Tuệ Tâm Care; số liệu là kịch bản, không phải cam kết doanh thu.", fill="#667085", font=safe_font(18))
    img.save(path)


def save_gmv_allocation_chart(path: Path) -> None:
    parts = [
        ("KTV trực tiếp", int(KTV_SHARE * 100), "#1F4D78"),
        ("Trưởng đoàn", int(TEAM_LEAD_SHARE * 100), "#2E74B5"),
        ("Nền tảng", int(PLATFORM_SHARE * 100), "#B11927"),
        ("Giám đốc phân phối Quận", int(DISTRICT_DIRECTOR_SHARE * 100), "#7C8DA5"),
        ("Affiliate trực tiếp", int(DIRECT_AFFILIATE_SHARE * 100), "#B88922"),
    ]
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw_chart_header(
        draw,
        "Phân bổ mục tiêu GMV Business 2 triệu USD/năm",
        f"65% cho đội trực tiếp triển khai; 35% cho nền tảng và kênh phân phối. Quy đổi hoạch định {int(FX):,} VND/USD = {float(BUSINESS_GMV_VND/Decimal('1000000000')):.1f} tỷ đồng/năm.".replace(",", ".").replace(".1f", ""),
    )
    x0, y0, total_w, h = 92, 260, 1416, 120
    cursor = x0
    label_font = safe_font(20, bold=True)
    value_font = safe_font(22, bold=True)
    for label, pct, color in parts:
        width = int(total_w * pct / 100)
        draw.rectangle((cursor, y0, cursor + width, y0 + h), fill=color)
        pct_text = f"{pct}%"
        bbox = draw.textbbox((0, 0), pct_text, font=value_font)
        draw.text((cursor + width / 2 - (bbox[2] - bbox[0]) / 2, y0 + 43), pct_text, fill="white", font=value_font)
        cursor += width

    group_font = safe_font(19, bold=True)
    first_group_w = int(total_w * 65 / 100)
    draw.text((x0, 405), "ĐỘI TRỰC TIẾP TRIỂN KHAI · 65%", fill="#1F4D78", font=group_font)
    draw.text((x0 + first_group_w + 18, 405), "NỀN TẢNG & PHÂN PHỐI · 35%", fill="#B11927", font=group_font)

    y = 465
    for idx, (label, pct, color) in enumerate(parts):
        col = idx % 3
        row = idx // 3
        x = 90 + col * 505
        yy = y + row * 120
        draw.rounded_rectangle((x, yy, x + 26, yy + 26), radius=4, fill=color)
        amount = BUSINESS_GMV_VND * Decimal(pct) / Decimal(100)
        draw.text((x + 42, yy - 3), f"{label}: {pct}%", fill="#1F2937", font=label_font)
        draw.text((x + 42, yy + 31), fmt_billion(amount, 2), fill="#667085", font=safe_font(19))
    draw.text((92, 812), "KTV nhận 60% theo ca/lượt thực hiện; trưởng đoàn nhận thêm 5% GMV cho trách nhiệm điều phối. Tip tự nguyện nằm ngoài GMV.", fill="#667085", font=safe_font(18))
    img.save(path)


def save_business_demand_chart(path: Path) -> None:
    spends = [("300.000đ/người/tháng", 300_000), ("600.000đ/người/tháng", 600_000), ("1.000.000đ/người/tháng", 1_000_000)]
    monthly_gmv = float(BUSINESS_GMV_VND / Decimal(12))
    counts = [monthly_gmv / spend for _, spend in spends]
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw_chart_header(
        draw,
        "Số người dùng trả tiền cần duy trì mỗi tháng",
        "Mục tiêu GMV Business 52,2 tỷ đồng/năm; số lượng giảm khi mức chi tiêu bình quân/người tăng.",
    )
    colors = ["#1F4D78", "#2E74B5", "#B88922"]
    max_value = max(counts)
    x0, plot_w, y0 = 530, 920, 230
    label_font = safe_font(24, bold=True)
    value_font = safe_font(25, bold=True)
    for idx, ((label, _), count) in enumerate(zip(spends, counts)):
        y = y0 + idx * 185
        draw.text((70, y + 30), label, fill="#1F2937", font=label_font)
        width = int(plot_w * count / max_value)
        draw.rounded_rectangle((x0, y, x0 + width, y + 96), radius=12, fill=colors[idx])
        text = f"{count:,.0f} người/tháng".replace(",", ".")
        bbox = draw.textbbox((0, 0), text, font=value_font)
        tx = x0 + width - (bbox[2] - bbox[0]) - 22
        draw.text((tx, y + 30), text, fill="white", font=value_font)
    draw.text((70, 812), "Đây là người dùng trả tiền hoạt động mỗi tháng, không phải số tài khoản đăng ký hoặc người nhận truyền thông.", fill="#667085", font=safe_font(18))
    img.save(path)


def save_break_even_chart(path: Path) -> None:
    fixed_levels = [50, 65, 80]
    variable_rates = [25, 30, 35]
    values: list[list[float]] = []
    for vr in variable_rates:
        row = []
        for fixed in fixed_levels:
            annual = float(FIXED_SITE_TARGET) * (1 - vr / 100) - fixed * 1_000_000 * 2 * 12
            row.append(annual / 1_000_000)
        values.append(row)
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw_chart_header(
        draw,
        "Độ nhạy lợi nhuận vận hành của hai cơ sở",
        "Đơn vị: triệu đồng/năm, trước thuế - lãi vay - khấu hao; mục tiêu doanh thu 2,5 tỷ đồng/năm.",
    )
    x0, y0, cell_w, cell_h = 430, 245, 315, 145
    header_font = safe_font(23, bold=True)
    value_font = safe_font(28, bold=True)
    draw.text((70, y0 - 70), "Chi phí biến đổi", fill="#667085", font=header_font)
    draw.text((x0 + 220, y0 - 70), "Chi phí cố định / cơ sở / tháng", fill="#667085", font=header_font)
    for col, fixed in enumerate(fixed_levels):
        text = f"{fixed} triệu"
        bbox = draw.textbbox((0, 0), text, font=header_font)
        draw.text((x0 + col * cell_w + cell_w / 2 - (bbox[2] - bbox[0]) / 2, y0 - 32), text, fill="#1F2937", font=header_font)
    for row, vr in enumerate(variable_rates):
        label = f"{vr}% doanh thu"
        draw.text((70, y0 + row * cell_h + 52), label, fill="#1F2937", font=header_font)
        for col, value in enumerate(values[row]):
            fill = "#EAF6EF" if value > 100 else ("#FFF5DA" if value >= 0 else "#FDEBEC")
            outline = "#18794E" if value > 100 else ("#B88922" if value >= 0 else "#B11927")
            x = x0 + col * cell_w
            y = y0 + row * cell_h
            draw.rounded_rectangle((x, y, x + cell_w - 18, y + cell_h - 18), radius=14, fill=fill, outline=outline, width=3)
            text = f"{value:+,.0f}".replace(",", ".")
            bbox = draw.textbbox((0, 0), text, font=value_font)
            draw.text((x + (cell_w - 18) / 2 - (bbox[2] - bbox[0]) / 2, y + 43), text, fill=outline, font=value_font)
    draw.text((72, 812), "Kết luận: 2,5 tỷ đồng doanh thu chỉ có lãi khi cơ cấu chi phí được giữ trong vùng kiểm soát; công suất đủ không đồng nghĩa lợi nhuận đủ.", fill="#667085", font=safe_font(18))
    img.save(path)


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "bottom", "start", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Column widths must total {CONTENT_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, width in (("w:tblW", CONTENT_DXA), ("w:tblInd", TABLE_INDENT_DXA)):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:type"), "dxa")
        node.set(qn("w:w"), str(width))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for idx, width in enumerate(widths_dxa):
        table.columns[idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, *, side: str = "left", color: str = BLUE, size: int = 16, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def shade_paragraph(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), color)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE, underline: bool = True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(r_fonts)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run("Trang ")
    set_run_font(label_run, size=8, color=MUTED)
    begin_run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_char1)
    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    instr_run._r.append(instr_text)
    separate_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(fld_sep)
    value_run = paragraph.add_run("1")
    set_run_font(value_run, size=8, color=MUTED)
    end_run = paragraph.add_run()
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char2)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, after: float = 6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_callout(doc: Document, title: str, text: str, *, kind: str = "blue") -> None:
    colors = {
        "blue": (PALE_BLUE, BLUE),
        "gold": (PALE_GOLD, BRAND_GOLD),
        "red": (PALE_RED, BRAND_RED),
        "green": ("EAF6EF", GREEN),
    }
    fill, border = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, fill)
    set_paragraph_border(p, color=border)
    r = p.add_run(title + "  ")
    set_run_font(r, bold=True, color=border)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, alignments: Sequence[str] | None = None, font_size: float = 9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row_idx, row_values in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(row_values):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                shade_cell(cell, "FAFAFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            alignment = (alignments[col_idx] if alignments else "left")
            p.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }[alignment]
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    apply_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, path: Path, caption: str, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.32))
    shape._inline.docPr.set("descr", alt_text)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(9)
    r = c.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, color=BLUE if level < 3 else NAVY, bold=True)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_document() -> None:
    fixed_chart = ASSET_DIR / "fixed-site-scenarios.png"
    gmv_chart = ASSET_DIR / "business-gmv-allocation.png"
    demand_chart = ASSET_DIR / "business-active-payers.png"
    breakeven_chart = ASSET_DIR / "fixed-site-break-even.png"
    save_fixed_site_chart(fixed_chart)
    save_gmv_allocation_chart(gmv_chart)
    save_business_demand_chart(demand_chart)
    save_break_even_chart(breakeven_chart)
    if LOGO_PATH.exists():
        with Image.open(LOGO_PATH) as logo:
            logo = logo.convert("RGBA")
            logo.thumbnail((512, 512), Image.Resampling.LANCZOS)
            logo.save(LOGO_SMALL_PATH, optimize=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, NAVY, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(3)
    hr = hp.add_run("TUỆ TÂM CARE  |  NGHIÊN CỨU THỊ TRƯỜNG & KẾ HOẠCH KINH DOANH")
    set_run_font(hr, size=8, color=MUTED, bold=True)
    set_paragraph_border(hp, side="bottom", color="E2E8F0", size=4, space=4)

    fp = section.footer.paragraphs[0]
    fr = fp.add_run("Tài liệu hoạch định nội bộ  |  Phiên bản 2.0  |  ")
    set_run_font(fr, size=8, color=MUTED)
    add_page_number(fp)

    # Cover: editorial_cover pattern on top of the standard business brief preset.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)
    if LOGO_SMALL_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        logo = p.add_run().add_picture(str(LOGO_SMALL_PATH), width=Inches(0.78))
        logo._inline.docPr.set("descr", "Biểu trưng Tuệ Tâm Care")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("TUỆ TÂM CARE")
    set_run_font(r, size=11, color=BRAND_RED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("NGHIÊN CỨU THỊ TRƯỜNG\nVÀ KẾ HOẠCH KINH DOANH")
    set_run_font(r, size=27, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Hai cơ sở chăm sóc tại Hà Nội & nền tảng Tuệ Tâm Business")
    set_run_font(r, size=15, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(
        "Cơ sở 1: Số 1 Hoàng Quán Chi, Dịch Vọng, Cầu Giấy\n"
        "Cơ sở 2: A11 LK6D BCA, Nguyễn Văn Lộc, Hà Đông"
    )
    set_run_font(r, size=11, color=MUTED)
    add_callout(
        doc,
        "Mục tiêu hoạch định",
        "Kiểm tra tính khả thi của mục tiêu 2,5 tỷ đồng doanh thu/năm tại hai cơ sở và 2 triệu USD GMV Business/năm, đồng thời xác định điều kiện vận hành để tăng trưởng có kiểm soát.",
        kind="gold",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("Phiên bản 2.1  •  26/07/2026\nCập nhật cấu trúc phân bổ GMV và định vị Tuệ Tâm Business")
    set_run_font(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu hoạch định - không phải cam kết doanh thu, tư vấn đầu tư, thuế hoặc pháp lý.")
    set_run_font(r, size=9, color=BRAND_RED, italic=True)

    page_break(doc)
    add_heading(doc, "Mục lục định hướng", 1)
    toc_items = [
        "1. Tóm tắt điều hành và quyết định đề xuất",
        "2. Phạm vi, định nghĩa và phương pháp",
        "3. Bức tranh thị trường Hà Nội - TP.HCM",
        "4. Phân tích hai điểm kinh doanh tại Hà Nội",
        "5. Kế hoạch doanh thu hai cơ sở cố định",
        "6. Quy mô và khả thi của mục tiêu Tuệ Tâm Business",
        "7. Cấu trúc phân bổ GMV - thu nhập đội ngũ - kênh phân phối",
        "8. Năng lực KTV và kiến trúc vận hành",
        "9. Hiệu quả tài chính, điểm hòa vốn và hoàn vốn",
        "10. Chiến lược thị trường và lộ trình 12 tháng",
        "11. KPI quản trị và cổng quyết định",
        "12. Rủi ro, pháp lý, dữ liệu và kiểm soát",
        "13. Kết luận tính khả thi",
        "Phụ lục: Công thức, bảng giả định và nguồn tham khảo",
    ]
    add_numbers(doc, toc_items)
    add_callout(
        doc,
        "Cách đọc nhanh",
        "Đọc phần 1 để nắm quyết định, phần 5 và 6 để hiểu sức nặng của hai mục tiêu doanh thu, sau đó dùng phần 9-12 làm điều kiện trước khi ký chi phí mở rộng.",
        kind="blue",
    )

    page_break(doc)
    add_heading(doc, "1. Tóm tắt điều hành và quyết định đề xuất", 1)
    add_callout(
        doc,
        "Kết luận trung tâm",
        "Mục tiêu 2,5 tỷ đồng/năm tại hai cơ sở là khả thi về công suất nhưng lợi nhuận còn phụ thuộc mạnh vào tiền thuê, lương cố định và tỷ lệ chia KTV. Mục tiêu Business 2 triệu USD/năm có thị trường đủ rộng nhưng là mục tiêu tăng trưởng quy mô mạng lưới; 16 KTV hiện tại không thể tự thực hiện toàn bộ khối lượng.",
        kind="red",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Kết quả hoạch định", "Ý nghĩa điều hành"],
        [
            ["Doanh thu hai cơ sở", "2,5 tỷ đ/năm", "208,3 triệu đ/tháng toàn hệ thống"],
            ["Nhịp bán cơ sở", "Khoảng 10 lượt/ngày/cơ sở", "Vé bình quân 350.000đ; 360 ngày/năm"],
            ["Khách thường xuyên", "208-694 người", "Tương ứng chi 1.000.000đ-300.000đ/người/tháng"],
            ["GMV Business", "2 triệu USD = 52,2 tỷ đ/năm", "4,35 tỷ đ/tháng ở trạng thái ổn định"],
            ["Doanh thu nền tảng", "10,44 tỷ đ/năm", "Nếu nền tảng ghi nhận 20% phí và đóng vai trò đại lý"],
            ["Đội trực tiếp triển khai", "33,93 tỷ đ/năm", "65% GMV: KTV 60% + trưởng đoàn 5%"],
            ["Kênh phân phối", "7,83 tỷ đ/năm", "15% GMV: Giám đốc phân phối Quận 5% + Affiliate trực tiếp 10%"],
            ["Người trả tiền Business", "4.350-14.500 người/tháng", "Tùy mức chi 1.000.000đ-300.000đ/người/tháng"],
        ],
        [2400, 2600, 4360],
        alignments=["left", "center", "left"],
    )
    add_heading(doc, "Quyết định đề xuất", 2)
    add_numbers(
        doc,
        [
            "Thông qua mục tiêu 2,5 tỷ đồng/năm cho hai cơ sở như một mục tiêu cơ sở, nhưng chỉ coi là thành công khi đạt đồng thời lợi nhuận vận hành và tỷ lệ quay lại.",
            "Định nghĩa 2 triệu USD là GMV mục tiêu ở trạng thái ổn định. Nếu khởi động từ số 0, nên đặt năm đầu đạt khoảng 1 triệu USD GMV và kết thúc năm ở nhịp 4,35 tỷ đồng/tháng; năm thứ hai mới thu đủ 2 triệu USD trong 12 tháng.",
            "Tách mạng lưới Business khỏi công suất 16 KTV tại cơ sở: xây đội KTV Business trực tiếp hưởng 60% GMV, trưởng đoàn hưởng 5% GMV, có tiêu chuẩn đào tạo, QR định danh và SLA chất lượng.",
            "Chỉ mở rộng TP.HCM theo mô hình pod sau khi Hà Nội chứng minh được tỷ lệ mua lại/gia hạn, biên đóng góp và kiểm soát khiếu nại; chưa nên mở cơ sở sở hữu mới chỉ để đuổi theo GMV.",
        ],
    )
    add_callout(
        doc,
        "Điều kiện không được bỏ qua",
        "GMV 52,2 tỷ đồng không phải doanh thu kế toán của Tuệ Tâm. Nền tảng nhận 20% tương đương 10,44 tỷ đồng; 65% là thu nhập của KTV và trưởng đoàn trực tiếp triển khai, 15% dành cho Giám đốc phân phối cấp Quận và Affiliate trực tiếp.",
        kind="gold",
    )

    add_heading(doc, "2. Phạm vi, định nghĩa và phương pháp", 1)
    add_heading(doc, "2.1. Phạm vi quyết định", 2)
    add_body(doc, "Báo cáo phục vụ chủ doanh nghiệp Tuệ Tâm Care trong việc đặt mục tiêu doanh thu, thiết kế mạng lưới bán Business, phân bổ hoa hồng, xác định công suất KTV và xây cổng quyết định trước khi mở rộng Hà Nội - TP.HCM.")
    add_heading(doc, "2.2. Những định nghĩa phải thống nhất", 2)
    add_table(
        doc,
        ["Khái niệm", "Định nghĩa dùng trong báo cáo"],
        [
            ["GMV Business", "Tổng giá trị giao dịch của sản phẩm dịch vụ Business đã thu tiền/đủ điều kiện ghi nhận; chưa trừ phân bổ cho KTV, trưởng đoàn, nền tảng và kênh giới thiệu."],
            ["Doanh thu nền tảng", "Phần phí 20% thuộc Tuệ Tâm nếu pháp nhân đóng vai trò đại lý/kết nối; cách hạch toán cuối cùng cần kế toán xác nhận."],
            ["Doanh thu cơ sở", "Giá trị dịch vụ tại hai cơ sở do Tuệ Tâm trực tiếp vận hành, không gồm tiền tip KTV."],
            ["Tip KTV", "Khoản khách tự nguyện trả ngoài bill dịch vụ, không đưa vào mục tiêu 2,5 tỷ đồng và không dùng để bù giá dịch vụ."],
            ["Người dùng trả tiền", "Người phát sinh thanh toán hợp lệ trong tháng; không đồng nghĩa tài khoản đăng ký hoặc người được tiếp cận."],
        ],
        [2300, 7060],
    )
    add_heading(doc, "2.3. Phương pháp", 2)
    add_bullets(
        doc,
        [
            "Bottom-up cho cơ sở: lượt/ngày × vé bình quân × ngày vận hành × số cơ sở.",
            "Bottom-up cho Business: người sử dụng trả tiền × chi tiêu/tháng; kiểm tra chéo bằng đơn hàng/hợp đồng, điểm triển khai và năng lực KTV.",
            "Độ nhạy tài chính: chi phí biến đổi 25%-35% doanh thu và chi phí cố định 50-80 triệu đồng/cơ sở/tháng.",
            "Nguồn công khai dùng để xác nhận độ lớn của tệp doanh nghiệp, lao động, văn phòng và bối cảnh wellness; số dự báo nội bộ luôn được gắn nhãn giả định.",
        ],
    )
    add_source_note(doc, "Nguồn chính: cấu hình dự án [S1], thống kê Hà Nội [S2][S3], TP.HCM [S4][S5], Savills [S6], GWI [S7].")
    add_callout(doc, "Mức độ tin cậy", "Cao với phép tính mục tiêu và cấu trúc phân bổ; trung bình với nhu cầu thị trường; thấp-trung bình với lợi nhuận vì chưa có chi phí thuê, lương, chia KTV, CAPEX và dữ liệu bán thật.", kind="blue")

    page_break(doc)
    add_heading(doc, "3. Bức tranh thị trường Hà Nội - TP.HCM", 1)
    add_heading(doc, "3.1. Nền cầu đủ lớn cho một mô hình ngách", 2)
    add_body(doc, "Hà Nội có 4,114 triệu người có việc làm tại thời điểm cuối quý I/2025, trong đó 1,977 triệu người làm việc ở khu vực thành thị [S2]. TP.HCM ghi nhận khoảng 4,73 triệu người có việc làm trong năm 2024 [S4]. Đây là nền dân số lao động lớn để tiếp cận các nhu cầu thư giãn, cổ-vai-gáy, phục hồi sau giờ làm và chăm sóc định kỳ.")
    add_body(doc, "Ở góc độ tài khoản tổ chức, Hà Nội có 217.414 doanh nghiệp đang hoạt động đến 31/05/2025 [S3]; TP.HCM có khoảng 345.000 doanh nghiệp đang hoạt động theo nguồn được ITPC dẫn năm 2025 [S5]. Mục tiêu Business không cần chiếm thị phần lớn của toàn thị trường, nhưng cần lọc đúng nơi có mật độ người dùng, không gian triển khai, nhu cầu chăm sóc tại chỗ và người có thẩm quyền mua hoặc giới thiệu dịch vụ.")
    add_body(doc, "Savills ghi nhận nguồn cung văn phòng Hà Nội Q1/2025 đạt 2,33 triệu m²; khu vực nội thành và phía Tây mỗi nơi chiếm khoảng 41% nguồn cung [S6]. Điều này ủng hộ chiến lược lấy Cầu Giấy làm cửa ngõ B2B phía Tây Hà Nội. Global Wellness Institute ước tính kinh tế wellness Việt Nam đạt 19,60 tỷ USD năm 2022 [S7], nhưng đây là thị trường rất rộng và không được dùng trực tiếp để suy ra doanh thu Tuệ Tâm.")
    add_heading(doc, "3.2. Nhu cầu doanh nghiệp: lý do mua và điều kiện mua", 2)
    add_bullets(
        doc,
        [
            "Lý do mua: sử dụng một sản phẩm dịch vụ chăm sóc tại chỗ tiện lợi vào buổi trưa hoặc khung giờ phù hợp; có thể mua cho cá nhân, nhóm, phòng ban, khách hàng, đối tác, sự kiện hoặc chương trình chăm sóc định kỳ.",
            "Người mua/người giới thiệu: CEO/chủ doanh nghiệp, kế toán, mua hàng, HR, Công đoàn, Office/Admin, ban quản lý tòa nhà, trưởng nhóm, đối tác hoặc bất kỳ cá nhân nào giới thiệu được khách hàng hợp lệ.",
            "Điều kiện mua: báo giá rõ theo đầu người/ca, chứng từ và hóa đơn, cam kết nhân sự, bảo mật thông tin, quy trình xử lý sự cố và báo cáo sau chương trình.",
            "Điều kiện mua lại/gia hạn: đúng giờ, chất lượng đồng đều, tỷ lệ sử dụng đủ cao, khiếu nại thấp và dữ liệu chứng minh người dùng thực sự sử dụng dịch vụ.",
        ],
    )
    add_source_note(doc, "WHO mô tả các vấn đề cơ xương khớp có thể làm giảm khả năng làm việc và tạo chi phí gián tiếp do vắng mặt hoặc giảm năng suất [S8]. ILO năm 2026 nhấn mạnh rủi ro tâm lý - xã hội và sức khỏe tinh thần tại nơi làm việc [S9]. Tuệ Tâm nên truyền thông là dịch vụ chăm sóc/thư giãn, không đưa ra tuyên bố chữa bệnh nếu chưa đủ căn cứ chuyên môn.")
    add_heading(doc, "3.3. Hà Nội trước, TP.HCM theo mô hình đối tác", 2)
    add_table(
        doc,
        ["Thị trường", "Lợi thế", "Rủi ro", "Chiến lược vào"],
        [
            ["Hà Nội", "Hai cơ sở làm hub; Cầu Giấy gần cụm văn phòng phía Tây; dễ kiểm soát chất lượng.", "Cạnh tranh cao; giao thông và khung giờ trưa; công suất KTV bị chia giữa tại điểm và Business.", "Bán sản phẩm dịch vụ trực tiếp + Affiliate minh bạch; pilot 10-20 điểm triển khai; xây playbook trước khi nhân rộng."],
            ["TP.HCM", "Tệp doanh nghiệp và lao động lớn; nhu cầu dịch vụ chăm sóc tại chỗ đa dạng.", "Khoảng cách địa lý; chi phí giám sát; chưa có hub chất lượng và đội triển khai riêng.", "Đối tác/pod trước; chỉ mở cơ sở sở hữu khi đạt ngưỡng GMV, mua lại và chất lượng."],
        ],
        [1400, 2800, 2460, 2700],
        font_size=8.8,
    )

    add_heading(doc, "4. Phân tích hai điểm kinh doanh tại Hà Nội", 1)
    add_heading(doc, "4.1. Cơ sở 1 - Số 1 Hoàng Quán Chi, Dịch Vọng, Cầu Giấy", 2)
    add_body(doc, "Vai trò chiến lược: hub thương hiệu, chất lượng và bán B2B cho khu vực Cầu Giấy - Duy Tân - Trần Thái Tông - Mỹ Đình. Nguồn cung văn phòng phía Tây Hà Nội lớn [S6], đồng thời Cầu Giấy nằm trong nhóm địa bàn có số doanh nghiệp đăng ký mới cao [S3].")
    add_table(
        doc,
        ["Điểm mạnh", "Điểm cần kiểm soát"],
        [
            ["Tiếp cận nhân sự văn phòng, quản lý và chủ doanh nghiệp; phù hợp làm nơi trải nghiệm trước khi ký gói Business.", "Khách có nhiều lựa chọn thay thế; phải khác biệt bằng đặt lịch tức thì, KTV có hồ sơ, bill minh bạch và chăm sóc sau dịch vụ."],
            ["Có thể bán cả tại điểm buổi tối và Business buổi trưa; tối ưu một tệp khách theo hai hoàn cảnh sử dụng.", "Giờ trưa ngắn, tắc đường và bãi đỗ xe có thể làm giảm chuyển đổi; cần khảo sát thực địa và lịch tuyến."],
        ],
        [4680, 4680],
    )
    add_heading(doc, "4.2. Cơ sở 2 - A11 LK6D BCA, Nguyễn Văn Lộc, Hà Đông", 2)
    add_body(doc, "Vai trò chiến lược: hub khách địa phương, gia đình, khách mua thẻ/gói dài hạn và điểm phục vụ khu vực Hà Đông - Mỗ Lao - Văn Quán. Hà Đông dẫn đầu số doanh nghiệp đăng ký mới tại Hà Nội đến 31/05/2025 trong số liệu được Sở Tài chính công bố [S3], nhưng tệp Business cần lọc theo mật độ người dùng, mức chi, điểm triển khai và khả năng mua lại dịch vụ.")
    add_table(
        doc,
        ["Điểm mạnh", "Điểm cần kiểm soát"],
        [
            ["Tệp cư dân ổn định thuận lợi cho gói định kỳ, referral gia đình và doanh nghiệp vừa/nhỏ quanh khu vực.", "Mức nhạy cảm giá có thể cao hơn Cầu Giấy; cần thiết kế gói theo giá trị, không giảm giá tràn lan."],
            ["Có thể làm hub vệ tinh cho tuyến Business phía Tây Nam Hà Nội.", "Không nên điều KTV chạy xa vào khu trung tâm giờ trưa nếu chi phí thời gian làm mất biên lợi nhuận."],
        ],
        [4680, 4680],
    )
    add_heading(doc, "4.3. Cấu hình công suất đã cập nhật", 2)
    add_table(
        doc,
        ["Cơ sở", "Giường gội", "Giường Foot", "Giường Body", "Tổng giường/ghế", "KTV"],
        [
            ["Cơ sở 1", "3", "6", "9", "18", "8"],
            ["Cơ sở 2", "3", "6", "9", "18", "8"],
            ["Toàn hệ thống", "6", "12", "18", "36", "16"],
        ],
        [2050, 1250, 1400, 1400, 1900, 1360],
        alignments=["left", "center", "center", "center", "center", "center"],
    )
    add_callout(doc, "Nút thắt thật", "Với 18 giường/ghế nhưng chỉ 8 KTV mỗi cơ sở, nhân sự mới là giới hạn trước vật lý. Hệ thống đặt lịch phải khóa đồng thời cả KTV và đúng loại giường theo toàn bộ thời lượng dịch vụ.", kind="gold")

    page_break(doc)
    add_heading(doc, "5. Kế hoạch doanh thu hai cơ sở cố định", 1)
    add_heading(doc, "5.1. Mục tiêu 2,5 tỷ đồng/năm tương đương điều gì?", 2)
    monthly_total = FIXED_SITE_TARGET / Decimal(12)
    monthly_branch = monthly_total / Decimal(2)
    add_table(
        doc,
        ["Đơn vị", "Mục tiêu", "Quy đổi vận hành"],
        [
            ["Toàn hệ thống", fmt_billion(FIXED_SITE_TARGET, 2) + "/năm", fmt_million(monthly_total, 1) + "/tháng"],
            ["Mỗi cơ sở", fmt_billion(FIXED_SITE_TARGET / 2, 2) + "/năm", fmt_million(monthly_branch, 1) + "/tháng"],
            ["Mỗi ngày/cơ sở", "-", fmt_million(monthly_branch / Decimal(30), 2) + "/ngày"],
        ],
        [2600, 2800, 3960],
        alignments=["left", "center", "center"],
    )
    add_figure(doc, fixed_chart, "Hình 1. Kịch bản doanh thu hai cơ sở cố định.", "Biểu đồ thanh ngang ba kịch bản doanh thu hai cơ sở Tuệ Tâm Care.")
    add_heading(doc, "5.2. Mức sử dụng cần đạt", 2)
    add_body(doc, "Kịch bản cơ sở 10 lượt/ngày/cơ sở với vé bình quân 350.000đ tạo 2,52 tỷ đồng/năm. Nếu một lượt trung bình 75 phút, mỗi cơ sở cần khoảng 12,5 giờ KTV/ngày. So với 8 KTV × 6 giờ phục vụ hữu ích = 48 giờ KTV/ngày, mức sử dụng lao động khoảng 26%. Về công suất, mục tiêu là khả thi và còn vùng đệm.")
    add_body(doc, "Tuy nhiên, công suất tổng có thể che giấu nghẽn cục bộ: Foot bận trong khi Body trống; tất cả khách tập trung 18h-21h; hoặc một KTV được chọn nhiều hơn. Vì vậy cần quản trị theo loại giường, KTV, khung giờ và dịch vụ, không chỉ nhìn phần trăm lấp đầy chung.")
    add_heading(doc, "5.3. Số khách thường xuyên cần xây", 2)
    add_table(
        doc,
        ["Chi tiêu bình quân/người/tháng", "Khách hoạt động cần có", "Hàm ý"],
        [
            ["300.000đ", "Khoảng 694 người", "Tệp rộng, tần suất thấp; cần CRM và nhắc lại tốt."],
            ["600.000đ", "Khoảng 347 người", "Kịch bản cơ sở cho khách định kỳ."],
            ["1.000.000đ", "Khoảng 208 người", "Tệp giá trị cao; phụ thuộc gói/thẻ và trải nghiệm nhất quán."],
        ],
        [3100, 2400, 3860],
        alignments=["center", "center", "left"],
    )
    add_heading(doc, "5.4. Thiết kế doanh thu theo tệp", 2)
    add_bullets(
        doc,
        [
            "Khách lẻ mới: tối ưu đặt lịch, cọc, check-in QR và trải nghiệm đầu tiên; mục tiêu chuyển thành khách quay lại trong 30 ngày.",
            "Khách định kỳ: gói 2-4 lượt/tháng, đổi lịch minh bạch, lịch sử sử dụng và nhắc lịch cá nhân hóa.",
            "Khách doanh nghiệp/đối tác: dùng cơ sở như điểm trải nghiệm và phục vụ ngoài giờ của gói Business.",
            "Affiliate: nhận thưởng theo bill dịch vụ đã hoàn tất và thu tiền; không thưởng cho việc tạo tài khoản hoặc tuyển người vào mạng lưới.",
        ],
    )

    page_break(doc)
    add_heading(doc, "6. Quy mô và khả thi của mục tiêu Tuệ Tâm Business", 1)
    add_heading(doc, "6.1. Đây là sản phẩm dịch vụ tại chỗ, không chỉ là phúc lợi doanh nghiệp", 2)
    add_body(doc, "Tuệ Tâm Business được định vị là sản phẩm dịch vụ chăm sóc trực tiếp được mang tới nơi có khách hàng: văn phòng, doanh nghiệp, tòa nhà, sự kiện, cuộc họp đối tác hoặc nhóm người dùng có nhu cầu. Doanh nghiệp có thể mua như một chương trình chăm sóc, nhưng đó chỉ là một trường hợp sử dụng; giá trị cốt lõi vẫn là dịch vụ thuận tiện, có KTV định danh, thời lượng rõ, chất lượng kiểm soát và lịch sử sử dụng minh bạch.")
    add_body(doc, "Người thanh toán có thể là doanh nghiệp, phòng ban, ban quản lý tòa nhà, nhà tài trợ, nhóm hoặc cá nhân. Người giới thiệu khách hàng hợp lệ có thể là CEO, kế toán, mua hàng, HR, Office/Admin, quản lý tòa nhà, đối tác hoặc cá nhân khác; cùng áp dụng một cơ chế Affiliate trực tiếp 10% khi nguồn khách được ghi nhận minh bạch và giao dịch đủ điều kiện.")
    add_callout(doc, "Nguyên tắc định vị", "Không bán một nhãn 'phúc lợi' duy nhất. Hãy bán một sản phẩm dịch vụ chăm sóc tại chỗ có thể đóng gói theo người, theo ca, theo đoàn, theo địa điểm hoặc theo ngân sách.", kind="blue")
    add_heading(doc, "6.2. Quy đổi mục tiêu", 2)
    add_body(doc, f"Tỷ giá hoạch định sử dụng là 26.100 VND/USD, gần tỷ giá mua chuyển khoản Vietcombank 26.110 ngày 20/07/2026 [S10]. Theo đó, 2 triệu USD GMV tương đương {fmt_billion(BUSINESS_GMV_VND, 1)} mỗi năm hoặc {fmt_billion(BUSINESS_GMV_VND/Decimal(12), 2)} mỗi tháng.")
    add_figure(doc, demand_chart, "Hình 2. Số người dùng trả tiền cần duy trì mỗi tháng để đạt GMV Business.", "Biểu đồ thanh số người dùng trả tiền cần thiết theo ba mức chi tiêu tháng.")
    add_heading(doc, "6.3. Kiểm tra chéo bằng đơn hàng và hợp đồng", 2)
    add_table(
        doc,
        ["Giá trị hợp đồng bình quân/năm", "Số hợp đồng tương đương cần duy trì", "Bình quân hợp đồng mới/tháng nếu chu kỳ 12 tháng"],
        [
            ["250 triệu đ", "Khoảng 209", "17-18"],
            ["500 triệu đ", "Khoảng 104", "8-9"],
            ["1 tỷ đ", "Khoảng 52", "4-5"],
        ],
        [3300, 3100, 2960],
        alignments=["center", "center", "center"],
    )
    add_body(doc, "Một điểm triển khai 100 người, tỷ lệ sử dụng 30% và mức chi 600.000đ/người/tháng tạo khoảng 18 triệu đồng GMV/tháng. Để đạt 4,35 tỷ đồng/tháng cần khoảng 242 điểm tương đương. Nếu tỷ lệ sử dụng 50%, con số giảm còn khoảng 145; nếu 80%, còn khoảng 91. Điểm triển khai có thể là doanh nghiệp, tòa nhà, phòng ban, sự kiện hoặc cộng đồng khách hàng; vì vậy tăng tỷ lệ sử dụng và mua lại quan trọng không kém số hợp đồng.")
    add_heading(doc, "6.4. Hai cách hiểu của mục tiêu 2 triệu USD", 2)
    add_table(
        doc,
        ["Cách hiểu", "Yêu cầu", "Đánh giá"],
        [
            ["Run-rate cuối năm", "Tháng 12 đạt 4,35 tỷ đ GMV; năm đầu thu khoảng 26,1 tỷ đ nếu tăng tuyến tính.", "Khả thi hơn; phù hợp giai đoạn xây mạng lưới."],
            ["Thu đủ trong 12 tháng đầu từ số 0", "Bình quân 4,35 tỷ đ/tháng ngay trong năm; tháng cuối có thể phải đạt 8-10 tỷ đ.", "Rất tham vọng; cần pipeline và mạng lưới KTV/trưởng đoàn có sẵn trước khi khởi động."],
            ["Trạng thái ổn định năm thứ hai", "Duy trì 4,35 tỷ đ/tháng trong 12 tháng.", "Khuyến nghị dùng làm định nghĩa mục tiêu chính thức."],
        ],
        [2500, 4200, 2660],
        font_size=8.9,
    )
    add_callout(doc, "Khuyến nghị mục tiêu", "Năm 1: 1 triệu USD GMV thu thật và kết thúc năm ở run-rate 2 triệu USD/năm. Năm 2: thu đủ 2 triệu USD, với điều kiện tỷ lệ mua lại/gia hạn và chất lượng dịch vụ vượt cổng quyết định.", kind="green")
    add_heading(doc, "6.5. Kiểm tra quy mô thị trường", 2)
    add_body(doc, "Hà Nội và TP.HCM có tổng cộng hơn 562 nghìn doanh nghiệp đang hoạt động theo hai nguồn công khai [S3][S5]. Mục tiêu 52,2 tỷ đồng/năm tương đương bình quân chưa tới 93.000đ/năm trên mỗi doanh nghiệp trong tập proxy này. Điều đó cho thấy giới hạn không nằm ở quy mô thị trường tổng, mà nằm ở khả năng tìm đúng doanh nghiệp, chốt hợp đồng, kích hoạt nhân viên và thực hiện dịch vụ đúng chuẩn.")
    add_body(doc, "Không nên dùng 19,60 tỷ USD của kinh tế wellness Việt Nam [S7] để nhân thị phần và gọi đó là dự báo doanh thu. Con số này bao gồm nhiều ngành không liên quan trực tiếp. Báo cáo chỉ dùng nó để xác nhận xu hướng rộng; mô hình doanh thu vẫn đi từ số người trả tiền và hợp đồng thực tế.")

    page_break(doc)
    add_heading(doc, "7. Cấu trúc phân bổ GMV - thu nhập đội ngũ - kênh phân phối", 1)
    add_body(doc, "Cấu trúc mới đặt người trực tiếp tạo ra dịch vụ ở vị trí trung tâm: 65% GMV thuộc đội triển khai, gồm 60% cho các KTV trực tiếp phục vụ và 5% cho trưởng đoàn chịu trách nhiệm điều phối. 35% còn lại gồm 20% cho nền tảng, 5% cho Giám đốc phân phối cấp Quận và 10% cho Affiliate trực tiếp giới thiệu khách hàng.")
    add_figure(doc, gmv_chart, "Hình 3. Phân bổ 100% GMV Business theo cấu trúc cập nhật.", "Thanh xếp chồng cho thấy 65 phần trăm thuộc đội trực tiếp triển khai và 35 phần trăm thuộc nền tảng cùng kênh phân phối.")
    add_heading(doc, "7.1. Phân bổ trên mỗi 100 triệu đồng GMV", 2)
    add_table(
        doc,
        ["Bên nhận", "Tỷ lệ", "Giá trị/100 triệu GMV", "Điều kiện chi trả"],
        [
            ["KTV trực tiếp thực hiện", "60%", "60 triệu đ", "Phân bổ cho đội KTV có check-in/out và thực hiện hợp lệ, theo thời lượng hoặc đơn vị dịch vụ đã chốt."],
            ["Trưởng đoàn Business", "5%", "5 triệu đ", "Khoản trách nhiệm điều phối trên tổng GMV ca/đoàn; có thể nhận thêm phần KTV nếu đồng thời trực tiếp phục vụ."],
            ["Nền tảng Tuệ Tâm", "20%", "20 triệu đ", "Vận hành công nghệ, tài khoản, thanh toán, QA, CSKH, thương hiệu, dự phòng và lợi nhuận."],
            ["Giám đốc phân phối cấp Quận", "5%", "5 triệu đ", "Phát triển địa bàn, quản trị pipeline, hỗ trợ điểm triển khai và chịu KPI chất lượng/phân phối."],
            ["Affiliate trực tiếp", "10%", "10 triệu đ", "Một nguồn giới thiệu được CRM ghi nhận; giao dịch đã thu tiền, hoàn tất và hết thời hạn hoàn/đối soát."],
        ],
        [2350, 1050, 2200, 3760],
        alignments=["left", "center", "center", "left"],
        font_size=8.8,
    )
    add_heading(doc, "7.2. Nguyên tắc phân bổ và đối soát", 2)
    add_numbers(
        doc,
        [
            "Khóa đúng 100% GMV: đội trực tiếp 65%, nền tảng 20%, Giám đốc phân phối cấp Quận 5% và Affiliate trực tiếp 10%; không phát sinh thêm tầng ngầm ngoài cấu trúc này.",
            "Quỹ KTV 60% được chia cho những KTV thực tế tham gia ca/đoàn theo đơn vị dịch vụ, thời lượng hoặc định mức đã công bố; tiền tip tự nguyện của khách nằm ngoài GMV và ngoài quỹ này.",
            "Trưởng đoàn nhận 5% tổng GMV vì trách nhiệm điều phối, điểm danh, chất lượng, thời gian, sự cố và đối soát. Nếu trưởng đoàn đồng thời trực tiếp phục vụ, phần KTV của họ phải được ghi riêng trong quỹ 60%.",
            "Affiliate trực tiếp chỉ có một nguồn hưởng 10% cho mỗi giao dịch. CEO, kế toán, mua hàng, HR, quản lý tòa nhà, đối tác hoặc cá nhân đều có thể là người giới thiệu nếu việc ghi nhận minh bạch và hợp lệ.",
            "Khi người giới thiệu là nhân sự hoặc người có quyền quyết định của bên mua, phải có quy trình công khai/xác nhận phù hợp của tổ chức để tránh xung đột lợi ích hoặc khoản chi cảm ơn không minh bạch.",
            "Mọi khoản phân bổ chỉ được ghi nhận sau giao dịch hợp lệ; áp dụng clawback khi hủy, hoàn tiền, gian lận, trùng nguồn hoặc không hoàn tất dịch vụ.",
        ],
    )
    add_heading(doc, "7.3. Gợi ý sử dụng 20% phí nền tảng", 2)
    add_table(
        doc,
        ["Hạng mục nội bộ", "Tỷ lệ trên GMV", "Giá trị/năm ở GMV mục tiêu"],
        [
            ["Công nghệ, dữ liệu, thanh toán", "3%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.03"), 2)],
            ["Sales Ops & quản trị tài khoản", "5%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.05"), 2)],
            ["QA, đào tạo, an toàn", "3%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.03"), 2)],
            ["CSKH & vận hành sự cố", "2%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.02"), 2)],
            ["Thương hiệu và nội dung", "2%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.02"), 2)],
            ["Dự phòng hoàn tiền/bad debt", "1%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.01"), 2)],
            ["Biên đóng góp mục tiêu", "4%", fmt_billion(BUSINESS_GMV_VND * Decimal("0.04"), 2)],
        ],
        [3900, 2300, 3160],
        alignments=["left", "center", "right"],
    )
    add_source_note(doc, "Bảng trên là ngân sách quản trị đề xuất, không phải số liệu thực tế. Thuế, phí thanh toán và cách ghi nhận doanh thu phải được kế toán/đơn vị tư vấn xác nhận.")

    add_heading(doc, "8. Năng lực KTV và kiến trúc vận hành", 1)
    add_heading(doc, "8.1. 16 KTV hiện tại phục vụ tốt mục tiêu cơ sở, không đủ cho GMV Business", 2)
    add_body(doc, "Nếu giá trị bình quân một lượt Business là 300.000đ, GMV 52,2 tỷ đồng tương đương khoảng 174.000 lượt/năm hoặc 483 lượt/ngày (360 ngày). Với 6 lượt/KTV/ngày, cần khoảng 81 KTV làm việc tương đương toàn thời gian; cộng 30% vùng đệm cho di chuyển, nghỉ và biến động cần khoảng 105 KTV hoạt động trong mạng lưới.")
    add_body(doc, "Nếu giá trị bình quân 500.000đ/lượt, nhu cầu còn khoảng 290 lượt/ngày, tương đương 48 KTV toàn thời gian hoặc khoảng 63 KTV sau vùng đệm. Nếu giữ gói micro-session từ 75.000đ/15 phút, khối lượng có thể lên gần 696.000 lượt/năm; khi đó cần khoảng 160 KTV mỗi ngày nếu mỗi người phục vụ 12 lượt ngắn, hoặc trên 200 KTV sau dự phòng. Đây là biến số phải khóa trước khi cam kết GMV.")
    add_body(doc, "Ở GMV mục tiêu, quỹ 60% dành cho KTV là 31,32 tỷ đồng/năm và quỹ 5% dành cho trưởng đoàn là 2,61 tỷ đồng/năm. Nếu quỹ KTV được phân bổ đều cho 105 KTV hoạt động, bình quân lý thuyết khoảng 24,9 triệu đồng/KTV/tháng; với mạng 63 KTV, khoảng 41,4 triệu đồng/KTV/tháng. Đây là phân bổ gộp trước thuế và chi phí cá nhân, không phải mức lương cam kết; thu nhập thực tế phải dựa trên ca/lượt hoàn tất của từng người.")
    add_callout(doc, "Ý nghĩa của tỷ lệ 65%", "Cấu trúc tạo động lực thu nhập mạnh cho người trực tiếp làm việc, nhưng nền tảng phải khóa định mức chia quỹ 60%, quy tắc trưởng đoàn 5% và lịch sử đối soát theo từng ca để tránh tranh chấp.", kind="gold")
    add_heading(doc, "8.2. Kiến trúc hub-and-pod đề xuất", 2)
    add_table(
        doc,
        ["Lớp vận hành", "Vai trò", "Kiểm soát bắt buộc"],
        [
            ["Hub cơ sở", "Đào tạo, trải nghiệm mẫu, dự phòng KTV, chăm sóc sau dịch vụ và xử lý sự cố.", "Chất lượng, vật tư, hồ sơ KTV, bill và dữ liệu khách."],
            ["Pod Business", "Đội triển khai theo cụm địa điểm, có trưởng đoàn và QR định danh; nhận tổng cộng 65% GMV.", "Check-in/out, GPS/địa điểm, thời lượng, đánh giá, chia quỹ KTV và đối soát."],
            ["Giám đốc phân phối cấp Quận", "Xây pipeline, phát triển điểm triển khai, phối hợp lịch và phản hồi chất lượng; nhận 5% GMV hợp lệ.", "SLA, nguồn CRM rõ, không tự sửa giá, không giữ tiền ngoài quy trình."],
            ["Nền tảng", "Booking, thanh toán, phân bổ, CRM, thông báo, báo cáo và cảnh báo rủi ro.", "Phân quyền, audit log, chống trùng hoa hồng, bảo mật dữ liệu."],
        ],
        [2100, 3900, 3360],
        font_size=8.9,
    )
    add_heading(doc, "8.3. Kỷ luật lịch và công suất", 2)
    add_bullets(
        doc,
        [
            "Một booking khóa đồng thời KTV, loại giường, cơ sở và toàn bộ thời lượng; ca 90 phút lúc 10h khóa đến 11h30.",
            "23h chỉ nhận ca 60 phút để kết thúc trước 24h; không nhận ca 90 phút lúc 23h.",
            "KTV đi Business phải được đánh dấu ngoài cơ sở trong toàn bộ thời gian di chuyển và phục vụ; không được hệ thống gán lịch tại quán chồng lên.",
            "Quản lý nhìn được số phút còn lại, ca kế tiếp, địa điểm và người chịu trách nhiệm để điều phối tức thời.",
        ],
    )

    page_break(doc)
    add_heading(doc, "9. Hiệu quả tài chính, điểm hòa vốn và hoàn vốn", 1)
    add_heading(doc, "9.1. Công suất đủ nhưng lợi nhuận chưa được chứng minh", 2)
    add_figure(doc, breakeven_chart, "Hình 4. Độ nhạy lợi nhuận vận hành của hai cơ sở.", "Ma trận độ nhạy lợi nhuận theo chi phí biến đổi và chi phí cố định.")
    add_body(doc, "Ở kịch bản chi phí biến đổi 30% doanh thu và chi phí cố định 65 triệu đồng/cơ sở/tháng, hai cơ sở tạo khoảng 190 triệu đồng lợi nhuận vận hành/năm trước thuế, lãi vay và khấu hao - biên khoảng 7,6%. Nếu chi phí cố định tăng lên 80 triệu đồng/cơ sở/tháng, mô hình lỗ khoảng 170 triệu đồng/năm. Vì vậy, mục tiêu doanh thu 2,5 tỷ đồng hợp lý về bán hàng nhưng cần kỷ luật chi phí.")
    add_heading(doc, "9.2. Điểm hòa vốn theo chi phí cố định", 2)
    add_table(
        doc,
        ["Chi phí cố định/cơ sở/tháng", "Điểm hòa vốn doanh thu hai cơ sở/năm", "So với mục tiêu 2,5 tỷ"],
        [
            ["50 triệu đ", "Khoảng 1,71 tỷ đ", "Vùng an toàn tốt"],
            ["65 triệu đ", "Khoảng 2,23 tỷ đ", "Có lãi mỏng"],
            ["80 triệu đ", "Khoảng 2,74 tỷ đ", "Mục tiêu hiện tại chưa đủ"],
        ],
        [3400, 3500, 2460],
        alignments=["center", "center", "center"],
    )
    add_heading(doc, "9.3. Hoàn vốn", 2)
    add_body(doc, "Chưa thể đưa ra thời gian hoàn vốn chính thức vì thiếu CAPEX thực tế, tiền cọc/thi công, thiết bị, vốn lưu động và chi phí khai trương. Công thức quản trị nên dùng: Thời gian hoàn vốn = Tổng vốn đầu tư ban đầu / Dòng tiền tự do hàng năm. Không dùng doanh thu hoặc GMV làm mẫu số.")
    add_table(
        doc,
        ["CAPEX hai cơ sở (giả định)", "Dòng tiền 190 triệu/năm", "Dòng tiền 550 triệu/năm"],
        [
            ["1,2 tỷ đ", "6,3 năm", "2,2 năm"],
            ["1,8 tỷ đ", "9,5 năm", "3,3 năm"],
            ["2,4 tỷ đ", "12,6 năm", "4,4 năm"],
        ],
        [3200, 3080, 3080],
        alignments=["center", "center", "center"],
    )
    add_source_note(doc, "Các con số CAPEX và dòng tiền trong bảng là độ nhạy minh họa, không phải định giá hay dự báo đầu tư.")
    add_heading(doc, "9.4. Phân biệt báo cáo quản trị và báo cáo kế toán", 2)
    add_bullets(
        doc,
        [
            "Tổng dòng tiền kinh tế ở mục tiêu: 2,5 tỷ doanh thu cơ sở + 52,2 tỷ GMV Business = 54,7 tỷ đồng giá trị giao dịch, nhưng không được cộng toàn bộ thành doanh thu Tuệ Tâm.",
            "Nếu nền tảng là đại lý, doanh thu quản trị có thể là 2,5 tỷ + 10,44 tỷ phí nền tảng = 12,94 tỷ đồng trước điều chỉnh kế toán.",
            "Nếu nền tảng là bên cung cấp chính, cách ghi nhận gross/net, VAT, hóa đơn đầu ra và thanh toán đối tác có thể khác; cần kế toán trưởng/tư vấn thuế kết luận bằng văn bản.",
        ],
    )

    page_break(doc)
    add_heading(doc, "10. Chiến lược thị trường và lộ trình 12 tháng", 1)
    add_heading(doc, "10.1. Định vị thương mại", 2)
    add_callout(doc, "Lời hứa thương hiệu", "Dịch vụ chăm sóc dễ đặt, đúng người, đúng thời lượng, minh bạch thanh toán - tại cơ sở hoặc được mang tới đúng nơi khách hàng cần.", kind="gold")
    add_heading(doc, "10.2. Gói sản phẩm đề xuất", 2)
    add_table(
        doc,
        ["Nhóm", "Sản phẩm", "Mục tiêu"],
        [
            ["Cá nhân tại cơ sở", "60/90 phút; gói định kỳ 2-4 lượt/tháng; thẻ dài hạn dùng chung hai cơ sở.", "Tăng tần suất, giữ chân và doanh thu ổn định."],
            ["Business thử nghiệm", "15/20/30 phút tại văn phòng, tối thiểu theo ca/đoàn để đủ hiệu quả tuyến.", "Giảm rào cản ký lần đầu; đo tỷ lệ tham gia."],
            ["Business định kỳ", "Gói tháng/quý theo người dùng, ca, đoàn, địa điểm hoặc ngân sách dịch vụ.", "Mua lại, dự báo công suất và GMV."],
            ["Executive/đối tác", "Buổi chăm sóc riêng cho lãnh đạo, khách hàng, đối tác hoặc sự kiện.", "Vé cao hơn và mở rộng đơn hàng theo nhóm/địa điểm."],
        ],
        [1900, 4500, 2960],
        font_size=8.9,
    )
    add_heading(doc, "10.3. Lộ trình 12 tháng", 2)
    add_table(
        doc,
        ["Giai đoạn", "Mục tiêu chính", "Đầu ra bắt buộc", "Cổng đi tiếp"],
        [
            ["0-90 ngày", "Ổn định hai cơ sở; pilot Business tại 10-20 doanh nghiệp/điểm triển khai ở Hà Nội.", "Giá/gói, hợp đồng mẫu, SLA, QR KTV, đối soát 65%/35% và dashboard thật.", "Tỷ lệ hoàn tất >95%; khiếu nại nghiêm trọng = 0; có ít nhất 5 khách Business trả tiền."],
            ["Tháng 4-6", "Đạt khoảng 10 lượt/ngày/cơ sở; 20-40 khách/điểm Business hoạt động.", "Playbook bán, Affiliate attribution, đào tạo trưởng đoàn, P&L theo cơ sở.", "Biên đóng góp dương; mua lại/gia hạn pilot ≥60%; phân bổ đối soát đúng."],
            ["Tháng 7-9", "Bổ nhiệm Giám đốc phân phối cấp Quận tại Hà Nội; thử TP.HCM theo pod.", "Mạng KTV đạt chuẩn; lịch tuyến; kiểm tra pháp lý và bảo hiểm.", "Chất lượng TP.HCM ngang Hà Nội; CAC hoàn vốn trong 6 tháng."],
            ["Tháng 10-12", "Chạm run-rate 4,35 tỷ đ GMV/tháng; chuẩn bị năm 2.", "Forecast 90 ngày, năng lực dự phòng, đơn hàng/hợp đồng gia hạn và kế hoạch vốn.", "Không mở rộng nếu mua lại/gia hạn <70%, biên nền tảng <4% GMV hoặc SLA không đạt."],
        ],
        [1500, 2600, 3000, 2260],
        font_size=8.3,
    )
    add_heading(doc, "10.4. Nhịp GMV năm đầu khuyến nghị", 2)
    add_body(doc, "Nếu chọn mục tiêu run-rate 2 triệu USD/năm ở tháng 12, GMV tháng có thể tăng theo bốn nấc: 0,3-0,8 tỷ/tháng trong quý I; 1,2-2,0 tỷ trong quý II; 2,5-3,5 tỷ trong quý III; 3,8-4,35 tỷ trong quý IV. Tổng GMV năm đầu khoảng 24-28 tỷ đồng là hợp lý hơn 52,2 tỷ đồng khi bắt đầu từ số 0.")
    add_body(doc, "Muốn thu đủ 52,2 tỷ đồng ngay năm đầu, Tuệ Tâm phải có sẵn pipeline lớn, mạng lưới KTV trực tiếp hàng chục đến hàng trăm người, đội trưởng đoàn đủ năng lực và vốn lưu động đối soát. Đây nên là kịch bản tăng tốc có điều kiện, không phải ngân sách mặc định.")

    add_heading(doc, "11. KPI quản trị và cổng quyết định", 1)
    add_heading(doc, "11.1. Hai cơ sở cố định", 2)
    add_table(
        doc,
        ["Nhóm KPI", "Chỉ số", "Ngưỡng định hướng"],
        [
            ["Nhu cầu", "Booking/ngày/cơ sở; conversion; nguồn booking", "10 lượt/ngày/cơ sở ở kịch bản cơ sở"],
            ["Khách", "Tỷ lệ quay lại 30/60/90 ngày; active customers; chi tiêu/tháng", "Tăng đều; không phụ thuộc voucher"],
            ["Vận hành", "Tỷ lệ hoàn tất; no-show; đúng giờ; sử dụng KTV/giường", ">95% hoàn tất; theo dõi peak theo loại giường"],
            ["Tài chính", "Doanh thu, giá vé, biên đóng góp, chi phí cố định, dòng tiền", "P&L riêng từng cơ sở; không gộp che lỗ"],
            ["Chất lượng", "NPS/đánh giá, khiếu nại/1.000 lượt, tái phục vụ", "Không có sự cố nghiêm trọng; xu hướng cải thiện"],
        ],
        [1900, 4100, 3360],
        font_size=8.8,
    )
    add_heading(doc, "11.2. Tuệ Tâm Business", 2)
    add_table(
        doc,
        ["Tầng funnel", "KPI cốt lõi", "Câu hỏi điều hành"],
        [
            ["Pipeline", "Khách/điểm triển khai đủ điều kiện, giá trị pipeline, thời gian chốt", "Có đủ 3-4 lần mục tiêu GMV 90 ngày không?"],
            ["Bán", "Win rate, GMV/hợp đồng, CAC hoa hồng", "Kênh nào tạo hợp đồng có lợi nhuận?"],
            ["Kích hoạt", "Người dùng đủ điều kiện, người trả tiền, tỷ lệ sử dụng", "Đơn hàng/hợp đồng đã mua có thực sự được sử dụng?"],
            ["Triển khai", "Fill rate, đúng giờ, lượt/KTV/ngày, GMV/KTV, thời gian di chuyển", "Có đủ KTV và trưởng đoàn tại đúng quận, đúng giờ?"],
            ["Mua lại", "Mua lại/gia hạn 90/180/365 ngày, churn GMV, expansion", "Khách quay lại vì giá trị hay chỉ vì giảm giá?"],
            ["Tài chính", "GMV, nền tảng 20%, đội trực tiếp 65%, phân phối 15%, đối soát", "Mỗi 100 triệu GMV đã được phân bổ đúng đủ chưa?"],
        ],
        [1800, 3900, 3660],
        font_size=8.8,
    )
    add_heading(doc, "11.3. Cổng dừng/mở rộng", 2)
    add_bullets(
        doc,
        [
            "Không mở quận mới nếu 3 tháng liên tiếp biên nền tảng âm hoặc tỷ lệ mua lại/gia hạn dưới 60%.",
            "Không tuyển thêm affiliate nếu dữ liệu nguồn giao dịch, chống trùng và clawback chưa vận hành ổn định.",
            "Không cam kết GMV vượt năng lực nếu fill rate KTV dưới 95% hoặc khiếu nại chất lượng tăng.",
            "Không mở cơ sở sở hữu ở TP.HCM nếu mô hình đối tác chưa đạt SLA, doanh thu tối thiểu và payback dự kiến.",
        ],
    )

    page_break(doc)
    add_heading(doc, "12. Rủi ro, pháp lý, dữ liệu và kiểm soát", 1)
    add_table(
        doc,
        ["Rủi ro", "Mức ảnh hưởng", "Kiểm soát đề xuất"],
        [
            ["GMV bị hiểu nhầm là doanh thu", "Rất cao", "Định nghĩa kế toán gross/net; chart of accounts; đối soát theo vai trò và pháp nhân."],
            ["Chi giới thiệu thiếu minh bạch/xung đột lợi ích", "Rất cao", "Affiliate một tầng; một nguồn/giao dịch; hợp đồng, thuế, audit log và xác nhận phù hợp khi người giới thiệu thuộc bên mua."],
            ["Thiếu KTV Business", "Rất cao", "Forecast theo quận; pool KTV đạt chuẩn; trưởng đoàn; SLA, dự phòng và giới hạn bán."],
            ["Chất lượng không đồng đều", "Cao", "Đào tạo, checklist, QR định danh, đánh giá, mystery audit, tạm khóa KTV không đạt."],
            ["Rò rỉ dữ liệu khách/nhân viên", "Cao", "Consent, tối thiểu hóa dữ liệu, phân quyền, mã hóa, log truy cập, quy trình sự cố theo Nghị định 13 [S11]."],
            ["Tập trung vào một đầu mối lớn", "Cao", "Giới hạn tỷ trọng GMV theo khách/kênh; pipeline đa dạng; điều khoản chấm dứt và bàn giao dữ liệu."],
            ["Chi phí cố định vượt kế hoạch", "Cao", "P&L theo cơ sở; ngân sách tuần/tháng; cổng chi; cảnh báo hòa vốn."],
            ["Tuyên bố sức khỏe quá mức", "Trung bình-cao", "Thông điệp chăm sóc/thư giãn; chuyên gia rà soát; không quảng cáo chữa bệnh khi chưa đủ điều kiện."],
        ],
        [2500, 1500, 5360],
        font_size=8.5,
    )
    add_heading(doc, "12.1. Affiliate trực tiếp và xung đột lợi ích", 2)
    add_body(doc, "Affiliate trực tiếp được hưởng 10% khi là nguồn giới thiệu hợp lệ của giao dịch, không phân biệt họ là CEO, kế toán, mua hàng, HR, quản lý tòa nhà, đối tác hay cá nhân khác. Tuy nhiên, nếu người giới thiệu đang làm việc cho bên mua hoặc có quyền phê duyệt nhà cung cấp, khoản chi phải được công khai, có căn cứ hợp đồng/thuế và tuân thủ quy định nội bộ của tổ chức. Phương án an toàn hơn trong trường hợp xung đột là chuyển quyền lợi thành chiết khấu cho tổ chức, tín dụng dịch vụ hoặc chương trình giới thiệu được công bố công khai.")
    add_heading(doc, "12.2. Dữ liệu cá nhân", 2)
    add_body(doc, "Nền tảng xử lý họ tên, số điện thoại, lịch sử đặt lịch, thanh toán, địa điểm, đánh giá và có thể cả thông tin sức khỏe nhạy cảm. Nghị định 13/2023/NĐ-CP yêu cầu mục đích xử lý rõ, sự đồng ý phù hợp, giới hạn dữ liệu và biện pháp bảo vệ [S11]. Cần lập hồ sơ xử lý dữ liệu, chính sách lưu/xóa, phân quyền theo vai trò và quy trình thông báo sự cố.")
    add_heading(doc, "12.3. Danh sách tư liệu bắt buộc chủ dự án cần bổ sung", 2)
    add_numbers(
        doc,
        [
            "Chi phí thật theo cơ sở: thuê, cọc, điện nước, lễ tân/quản lý, lương cố định, chia KTV, vật tư, vệ sinh, phần mềm, khấu hao và thuế.",
            "CAPEX đã đầu tư: thi công, giường/ghế, thiết bị, biển bảng, PCCC/an ninh trật tự, vốn lưu động.",
            "Giá gói Business chính thức: mức tối thiểu mỗi ca/đoàn, giá theo 15/20/30 phút, phí di chuyển và điều kiện hủy.",
            "Quy chế phân bổ chính thức: KTV trực tiếp 60%, trưởng đoàn 5%, nền tảng 20%, Giám đốc phân phối cấp Quận 5% và Affiliate trực tiếp 10%; kèm công thức chia quỹ KTV theo ca/lượt.",
            "Dữ liệu bán thật 8-12 tuần: lead, booking, bill, dịch vụ, KTV, khung giờ, cơ sở, voucher, no-show, hoàn tiền, quay lại.",
            "Hợp đồng và ý kiến chuyên môn về pháp lý massage/xoa bóp, lao động/đối tác KTV, thuế, hóa đơn, dữ liệu cá nhân, chi giới thiệu và kiểm soát xung đột lợi ích.",
        ],
    )

    add_heading(doc, "13. Kết luận tính khả thi", 1)
    add_table(
        doc,
        ["Hạng mục", "Điểm khả thi", "Kết luận"],
        [
            ["Doanh thu 2,5 tỷ tại hai cơ sở", "4/5 về công suất", "Khả thi với khoảng 10 lượt/ngày/cơ sở; lợi nhuận phụ thuộc chi phí."],
            ["Lợi nhuận hai cơ sở", "2,5/5 hiện tại", "Chưa đủ dữ liệu chi phí; có thể từ lỗ 170 triệu đến lãi 550 triệu/năm trong dải giả định."],
            ["GMV Business 2 triệu USD", "3/5 ở trạng thái ổn định", "Thị trường đủ lớn nhưng cần mạng lưới 60-100+ KTV tương đương; cao hơn nhiều nếu bán micro-session 75.000đ."],
            ["Cấu trúc phân bổ 65%/35%", "3,5/5", "Rõ động lực KTV; cần định mức chia quỹ 60%, trách nhiệm trưởng đoàn, attribution, clawback, thuế và kiểm soát xung đột lợi ích."],
            ["Mở rộng TP.HCM", "2,5/5 trước pilot", "Nên vào bằng đối tác/pod; chưa nên đầu tư hub sở hữu trước khi chứng minh unit economics."],
        ],
        [2700, 2200, 4460],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Quyết định cuối",
        "Dự án đáng để triển khai theo từng cổng kiểm chứng. Hai cơ sở là nền tảng vận hành và thương hiệu; Business là động cơ tăng trưởng. Không nên dùng GMV để che lỗ cơ sở, dùng tiền tip để bù giá dịch vụ, hoặc bán vượt năng lực triển khai của đội KTV.",
        kind="green",
    )
    add_heading(doc, "Việc cần làm ngay trong 30 ngày", 2)
    add_numbers(
        doc,
        [
            "Khóa P&L thật của từng cơ sở và xác nhận điểm hòa vốn theo tháng.",
            "Chuẩn hóa một gói Business pilot, hợp đồng, báo giá, SLA và quy trình QR từ bắt đầu đến thanh toán.",
            "Tạo CRM attribution cho hai vai trò phân phối: Giám đốc phân phối cấp Quận và Affiliate trực tiếp; đồng thời tạo sổ phân bổ riêng cho KTV và trưởng đoàn theo từng ca.",
            "Tuyển/kiểm định pool KTV Business theo quận; giới hạn số hợp đồng theo năng lực đã xác nhận.",
            "Chạy 10 điểm pilot trả phí thuộc nhiều tình huống sử dụng (doanh nghiệp, tòa nhà, phòng ban, sự kiện/nhóm) và đo tỷ lệ sử dụng, NPS, khiếu nại, GMV/người, GMV/KTV, tỷ lệ mua lại.",
        ],
    )

    page_break(doc)
    add_heading(doc, "Phụ lục A. Công thức và bảng giả định", 1)
    add_heading(doc, "A1. Hai cơ sở", 2)
    add_table(
        doc,
        ["Công thức", "Giá trị"],
        [
            ["Doanh thu năm", "Lượt/ngày/cơ sở × vé bình quân × 360 ngày × 2 cơ sở"],
            ["Kịch bản cơ sở", "10 × 350.000đ × 360 × 2 = 2,52 tỷ đ"],
            ["Giờ KTV cần/ngày/cơ sở", "10 lượt × 1,25 giờ = 12,5 giờ"],
            ["Công suất KTV giả định", "8 KTV × 6 giờ hữu ích = 48 giờ/ngày"],
            ["Mức sử dụng KTV", "12,5 / 48 ≈ 26%"],
        ],
        [3700, 5660],
    )
    add_heading(doc, "A2. Business", 2)
    add_table(
        doc,
        ["Công thức", "Giá trị"],
        [
            ["GMV VND", "2.000.000 USD × 26.100 = 52,2 tỷ đ"],
            ["GMV tháng", "52,2 / 12 = 4,35 tỷ đ"],
            ["Doanh thu nền tảng", "52,2 × 20% = 10,44 tỷ đ"],
            ["KTV trực tiếp", "52,2 × 60% = 31,32 tỷ đ"],
            ["Trưởng đoàn", "52,2 × 5% = 2,61 tỷ đ"],
            ["Giám đốc phân phối cấp Quận", "52,2 × 5% = 2,61 tỷ đ"],
            ["Affiliate trực tiếp", "52,2 × 10% = 5,22 tỷ đ"],
        ],
        [3700, 5660],
    )
    add_heading(doc, "A3. Giả định cần thay bằng dữ liệu thật", 2)
    add_table(
        doc,
        ["Giả định", "Dải dùng trong báo cáo", "Ưu tiên xác minh"],
        [
            ["Vé bình quân cơ sở", "300.000-400.000đ", "Cao"],
            ["Thời lượng bình quân", "75 phút", "Cao"],
            ["Giờ phục vụ hữu ích/KTV/ngày", "6 giờ", "Cao"],
            ["Chi phí biến đổi cơ sở", "25%-35% doanh thu", "Rất cao"],
            ["Chi phí cố định/cơ sở/tháng", "50-80 triệu đ", "Rất cao"],
            ["Chi tiêu Business/người/tháng", "300.000-1.000.000đ", "Cao"],
            ["Vùng đệm nhân lực Business", "30%", "Trung bình"],
        ],
        [3300, 3600, 2460],
        alignments=["left", "center", "center"],
    )

    add_heading(doc, "Phụ lục B. Danh mục nguồn tham khảo", 1)
    add_body(doc, "Các nguồn dưới đây dùng để xác nhận bối cảnh và đầu vào. Dữ liệu doanh thu, chi phí và chuyển đổi của Tuệ Tâm vẫn cần được thay bằng dữ liệu vận hành thật.")
    for source in SOURCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"[{source.code}] {source.title}. ")
        set_run_font(r, size=9.5, bold=True, color=INK)
        r = p.add_run(f"{source.publisher}, {source.date}. {source.use}")
        set_run_font(r, size=9.5, color=INK)
        if source.url:
            p.add_run(" ")
            add_hyperlink(p, "Mở nguồn", source.url)

    add_heading(doc, "Phụ lục C. Ghi chú kiểm định", 1)
    add_bullets(
        doc,
        [
            "Các tỷ lệ phân bổ GMV cộng đủ 100%: KTV 60% + trưởng đoàn 5% + nền tảng 20% + Giám đốc phân phối cấp Quận 5% + Affiliate trực tiếp 10%.",
            "Tiền tip không nằm trong bill, doanh thu cơ sở hoặc GMV dùng để tính hoa hồng.",
            "Mọi kịch bản tài chính chưa gồm thuế, lãi vay và khấu hao nếu không ghi rõ.",
            "Nguồn công khai có khác biệt về thời điểm và định nghĩa; báo cáo ghi rõ ngày và chỉ dùng cho phạm vi phù hợp.",
            "Kết luận tài chính được xếp 'chia sẻ có điều kiện' cho đến khi có P&L thật, CAPEX và dữ liệu bán 8-12 tuần.",
        ],
    )

    # Core properties and save.
    doc.core_properties.title = "Nghiên cứu thị trường và kế hoạch kinh doanh Tuệ Tâm Care v2.1"
    doc.core_properties.subject = "Hai cơ sở Hà Nội và mô hình Tuệ Tâm Business"
    doc.core_properties.author = "Tuệ Tâm Care"
    doc.core_properties.keywords = "Tuệ Tâm Care, nghiên cứu thị trường, kế hoạch kinh doanh, GMV, Business, Hà Nội, TP.HCM"
    doc.core_properties.comments = "Báo cáo hoạch định nội bộ được xây dựng ngày 26/07/2026."
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
