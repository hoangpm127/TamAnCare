from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\DELL\Desktop\Kyoto Masage")
SHOT = ROOT / "artifacts" / "tuetam-user-guide" / "screenshots"
OUT = ROOT / "artifacts" / "Tue_Tam_Care_Cam_Nang_Gia_Tri_Huong_Dan_Su_Dung_v1.0.docx"
LOGO = ROOT / "public" / "logo.png"

BURGUNDY = "9F1D20"
DARK_BROWN = "291714"
GOLD = "D8B46A"
GREEN = "16784A"
TEAL = "276C75"
CREAM = "FFF9F4"
PALE_RED = "FFF1EF"
PALE_GOLD = "FFF8E6"
PALE_GREEN = "EEF8F3"
PALE_BLUE = "EDF7FA"
MUTED = "665B55"
LIGHT_BORDER = "E7D8CD"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color=LIGHT_BORDER, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_keep(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_width(table, widths: Iterable[int]) -> None:
    widths = list(widths)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_BROWN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 30, DARK_BROWN, 0, 14),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 17, BURGUNDY, 18, 10),
        ("Heading 2", 13.5, DARK_BROWN, 14, 7),
        ("Heading 3", 11.5, TEAL, 10, 5),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            st.paragraph_format.page_break_before = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Guide Eyebrow" not in doc.styles:
        eyebrow = doc.styles.add_style("Guide Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
        eyebrow.font.name = "Calibri"
        eyebrow.font.size = Pt(8.5)
        eyebrow.font.bold = True
        eyebrow.font.color.rgb = RGBColor.from_string(GOLD)
        eyebrow.paragraph_format.space_after = Pt(5)
        eyebrow.paragraph_format.keep_with_next = True

    if "Guide Small" not in doc.styles:
        small = doc.styles.add_style("Guide Small", WD_STYLE_TYPE.PARAGRAPH)
        small.font.name = "Calibri"
        small.font.size = Pt(8.5)
        small.font.color.rgb = RGBColor.from_string(MUTED)
        small.paragraph_format.space_after = Pt(4)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.26)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.12


def add_header_footer(doc: Document) -> None:
    for sec in doc.sections:
        header = sec.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.7))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, [6600, 2760])
        set_repeat_table_header(table.rows[0])
        left, right = table.rows[0].cells
        for cell in (left, right):
            set_cell_margins(cell, 0, 0, 0, 0)
        p = left.paragraphs[0]
        if LOGO.exists():
            r = p.add_run()
            shape = r.add_picture(str(LOGO), width=Inches(0.26))
            set_alt_text(shape, "Logo Tuệ Tâm Care", "Biểu trưng Tuệ Tâm Care")
        r = p.add_run("  TUỆ TÂM CARE")
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(BURGUNDY)
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("CẨM NANG SỬ DỤNG")
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)

        footer = sec.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run("Tài liệu vận hành & đào tạo  •  Trang ")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        fr = p.add_run()
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(fr, "PAGE")


def add_callout(doc: Document, title: str, text: str, fill=PALE_GOLD, accent=GOLD) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "5")
        el.set(qn("w:color"), accent)
        borders.append(el)
    p_pr.append(borders)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r.font.size = Pt(10.5)
    r.add_break()
    body = p.add_run(text)
    body.font.color.rgb = RGBColor.from_string(DARK_BROWN)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def new_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract_id = getattr(doc, "_tuetam_decimal_abstract_id", None)
    if abstract_id is None:
        abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
        abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal")
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), "%1.")
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "420")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "420")
        ind.set(qn("w:hanging"), "240")
        p_pr.extend([tabs, ind])
        lvl.extend([start, fmt, text, jc, p_pr])
        abstract.append(lvl)
        # OOXML requires all abstractNum elements before all num instances.
        first_num_index = next(
            (idx for idx, child in enumerate(numbering) if child.tag == qn("w:num")),
            len(numbering),
        )
        numbering.insert(first_num_index, abstract)
        setattr(doc, "_tuetam_decimal_abstract_id", abstract_id)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def add_steps(doc: Document, items: Iterable[str]) -> None:
    num_id = new_decimal_numbering(doc)
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.29)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(4)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.append(num_pr)
        p.add_run(item)


def add_screenshot(doc: Document, filename: str, caption: str, alt: str, width=2.68) -> None:
    path = SHOT / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    r = p.add_run()
    shape = r.add_picture(str(path), width=Inches(width))
    set_alt_text(shape, caption, alt)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.keep_with_next = False


def add_feature_page(
    doc: Document,
    heading: str,
    value: str,
    steps: Iterable[str],
    result: str,
    image: str | None = None,
    caption: str | None = None,
    note: str | None = None,
) -> None:
    doc.add_heading(heading, level=2)
    add_callout(doc, "Giá trị tiện ích", value, PALE_GOLD, "A66A08")
    doc.add_heading("Cách sử dụng", level=3)
    add_steps(doc, steps)
    add_callout(doc, "Kết quả cần thấy", result, PALE_GREEN, GREEN)
    if note:
        p = doc.add_paragraph()
        r = p.add_run("Lưu ý: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BURGUNDY)
        p.add_run(note)
    if image and caption:
        add_screenshot(doc, image, caption, heading)


def add_role_table(doc: Document) -> None:
    rows = [
        ("Khách vãng lai", "Xem dịch vụ, ưu đãi, KTV; đặt lịch không cần đăng nhập", "Không thấy dữ liệu riêng của người khác"),
        ("Khách thành viên", "Đơn, voucher, QR check-in, lịch sử, thông báo", "Chỉ dữ liệu cá nhân"),
        ("Affiliate", "Mã/QR giới thiệu, chuyển đổi và thu nhập", "Chỉ hiệu quả giới thiệu của mình"),
        ("KTV", "Lịch cá nhân, điều phối, Business, QR, hồ sơ, thu nhập", "Chỉ ca được phân công và dữ liệu cần phục vụ"),
        ("Lễ tân", "Tạo khách, Booking, Lịch, Phòng, CRM tại quầy", "Chỉ cơ sở phụ trách; không được ghi chi"),
        ("Quản lý cơ sở", "Vận hành, tài chính, chi phí, QR, nhân sự trong cơ sở", "Chỉ cơ sở được giao"),
        ("Admin/Chủ", "Toàn hệ thống, cấu hình, báo cáo, cơ hội đầu tư", "Quyền cao nhất; bắt buộc kiểm soát MFA"),
        ("Nhà đầu tư", "Hiệu quả, hoàn vốn, cơ hội mới, đặc quyền, bản tin", "Chỉ xem phạm vi đã đầu tư; không sửa vận hành"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1900, 4500, 2960])
    hdr = table.rows[0].cells
    for idx, text in enumerate(("Vai trò", "Tính năng chính", "Giới hạn dữ liệu")):
        hdr[idx].text = text
        shade(hdr[idx], DARK_BROWN)
        cell_border(hdr[idx], DARK_BROWN)
        set_cell_margins(hdr[idx])
        for run in hdr[idx].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for i, row_data in enumerate(rows):
        cells = table.add_row().cells
        set_row_keep(table.rows[-1])
        for j, text in enumerate(row_data):
            cells[j].text = text
            shade(cells[j], "FFFFFF" if i % 2 == 0 else "FFF9F4")
            cell_border(cells[j])
            set_cell_margins(cells[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_flow_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1600, 2440, 2660, 2660])
    headers = ("Sự kiện", "Khách hàng", "KTV/Lễ tân/QL", "Admin & tài chính")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shade(cell, BURGUNDY)
        cell_border(cell, BURGUNDY)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for i, data in enumerate(rows):
        cells = table.add_row().cells
        set_row_keep(table.rows[-1])
        for j, text in enumerate(data):
            cells[j].text = text
            shade(cells[j], "FFFFFF" if i % 2 == 0 else CREAM)
            cell_border(cells[j])
            set_cell_margins(cells[j])


def page_break(doc: Document) -> None:
    # Feature content is allowed to flow naturally; major sections use the
    # Heading 1 style with page-break-before. This prevents orphan blank pages.
    return


def build() -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    props = doc.core_properties
    props.title = "Tuệ Tâm Care — Cẩm nang giá trị tiện ích & hướng dẫn sử dụng"
    props.subject = "Hướng dẫn chi tiết các luồng người dùng Tuệ Tâm Care"
    props.author = "Tuệ Tâm Care"
    props.keywords = "Tuệ Tâm Care, hướng dẫn sử dụng, vận hành, khách hàng, KTV, quản lý, admin, nhà đầu tư"

    # Cover: editorial_cover template.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    if LOGO.exists():
        shape = p.add_run().add_picture(str(LOGO), width=Inches(0.85))
        set_alt_text(shape, "Logo Tuệ Tâm Care", "Biểu trưng chính thức của Tuệ Tâm Care")
    p = doc.add_paragraph("TUỆ TÂM CARE", style="Guide Eyebrow")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("CẨM NANG GIÁ TRỊ TIỆN ÍCH\n& HƯỚNG DẪN SỬ DỤNG", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Từ trải nghiệm khách hàng đến vận hành, tài chính và đầu tư", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    shape = p.add_run().add_picture(str(ROOT / "public" / "tue-tam-hero.png"), width=Inches(5.75))
    set_alt_text(shape, "Không gian Tuệ Tâm Care", "Hình ảnh nhận diện không gian chăm sóc Tuệ Tâm Care")
    add_callout(
        doc,
        "Phiên bản hướng dẫn UAT · 07/2026",
        "Tài liệu dành cho đào tạo, trình diễn và chuẩn hóa thao tác. Ảnh minh họa được chụp trực tiếp từ bản Railway đang triển khai; số liệu UAT có thể thay đổi khi giao dịch mới phát sinh.",
        PALE_RED,
        BURGUNDY,
    )
    p = doc.add_paragraph("TÀI LIỆU VẬN HÀNH & GIỚI THIỆU GIẢI PHÁP", style="Guide Eyebrow")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_break(doc)

    doc.add_heading("Cách đọc tài liệu", level=1)
    add_callout(doc, "Mục tiêu", "Giúp bất kỳ người dùng nào hiểu nhanh giá trị của nền tảng, biết bấm ở đâu, kết quả nào là đúng và dữ liệu được chuyển tiếp cho vai trò liên quan như thế nào.", PALE_BLUE, TEAL)
    doc.add_heading("Mục lục định hướng", level=2)
    add_bullets(doc, [
        "Phần I — Giá trị tổng thể và bản đồ vai trò",
        "Phần II — Khách hàng, Affiliate và Tuệ Tâm Business",
        "Phần III — Kỹ thuật viên",
        "Phần IV — Lễ tân, Quản lý cơ sở và Admin/Chủ",
        "Phần V — Nhà đầu tư",
        "Phần VI — Logic liên vai trò: Booking, QR, tiền và thông báo",
        "Phần VII — Checklist vận hành thử và hỗ trợ sử dụng",
    ])
    doc.add_heading("Quy ước trong ảnh", level=2)
    add_bullets(doc, [
        "Màu đỏ rượu: hành động chính, điều hướng đang chọn hoặc trạng thái cần chú ý.",
        "Màu xanh: sẵn sàng, rảnh, xác nhận thành công hoặc trạng thái vận hành tích cực.",
        "Màu vàng: thông tin tài chính, đặt cọc, đặc quyền hoặc dữ liệu chờ đối soát.",
        "Chuông: thông báo theo vai trò; QR: check-in/check-out hoặc nhận diện nguồn phục vụ.",
    ])
    add_callout(doc, "Nguyên tắc an toàn", "Không chia sẻ tài khoản quản trị. Lễ tân không ghi nhận khoản chi; Nhà đầu tư chỉ xem; Quản lý cơ sở chỉ thao tác trong phạm vi được giao; Admin/Chủ nên bật MFA.", PALE_RED, BURGUNDY)
    page_break(doc)

    doc.add_heading("PHẦN I — GIÁ TRỊ TỔNG THỂ", level=1)
    doc.add_heading("1. Tuệ Tâm Care tạo ra giá trị gì?", level=2)
    add_bullets(doc, [
        "Với xã hội: tăng khả năng tiếp cận chăm sóc sức khỏe định kỳ, minh bạch dịch vụ và chuyên nghiệp hóa nghề KTV.",
        "Với khách hàng: đặt lịch nhanh, biết rõ chi phí, chọn KTV phù hợp, theo dõi dịch vụ và nhận thông báo mạch lạc.",
        "Với KTV: lịch làm việc rõ ràng, điều phối công bằng, hồ sơ năng lực được chuẩn hóa và thu nhập có lịch sử đối chiếu.",
        "Với cơ sở: giảm xung đột lịch–giường–KTV, số hóa CRM tại quầy, quản trị chi phí và công suất theo thời gian thực.",
        "Với Chủ doanh nghiệp: một nguồn dữ liệu xuyên suốt từ Booking đến doanh thu, chi phí, lợi nhuận, mở rộng Business và đầu tư.",
        "Với Nhà đầu tư: số liệu chỉ xem theo phạm vi vốn, báo cáo hiệu quả, dự kiến hoàn vốn và pipeline cơ hội tách biệt khỏi cơ sở đang vận hành.",
    ])
    doc.add_heading("2. Bản đồ vai trò và quyền dữ liệu", level=2)
    add_role_table(doc)
    page_break(doc)

    doc.add_heading("PHẦN II — KHÁCH HÀNG, AFFILIATE & BUSINESS", level=1)
    add_feature_page(
        doc,
        "1. Khám phá dịch vụ và tạo tài khoản",
        "Khách chưa đăng nhập vẫn xem được dịch vụ, ưu đãi và KTV. Tạo tài khoản giúp khôi phục Đơn của tôi, nhận voucher chào mừng và tự điền thông tin khi đặt lịch.",
        [
            "Mở đường dẫn Tuệ Tâm Care hoặc quét QR giới thiệu.",
            "Xem nhanh dịch vụ, ưu đãi, đội ngũ KTV và đánh giá trên Trang chủ.",
            "Chọn Tạo tài khoản để nhận quyền lợi lần đầu; điền đúng họ tên và số điện thoại có thể nhận thông báo.",
            "Sau khi đăng nhập, kiểm tra tab Tôi để xác nhận tên và ảnh đại diện; thông tin này sẽ được dùng mặc định khi đặt lịch.",
        ],
        "Trang chủ hiển thị đúng nhận diện; tài khoản thành viên xuất hiện trong mục Tôi; voucher đủ điều kiện được gợi ý ở bước đặt cọc.",
        "01-khach-hang-trang-chu.png",
        "Hình 1. Trang chủ khách hàng — điểm vào nhanh cho Đặt lịch, Ưu đãi, Affiliate và tài khoản.",
        "Không dùng chung số điện thoại cho nhiều khách nếu muốn lịch sử và ưu đãi được tính chính xác.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "2. Chọn hình thức đặt lịch",
        "Bốn hình thức giúp hệ thống hiểu ngữ cảnh phục vụ ngay từ đầu: Cá nhân, Mời bạn, Mời sếp và Tuệ Tâm Business.",
        [
            "Bấm dấu + hoặc Đặt lịch trên thanh điều hướng.",
            "Chọn Cá nhân khi tự sử dụng; Mời bạn khi đi cùng bạn bè; Mời sếp khi cần ghi nhận ngữ cảnh tiếp đón; Business khi phục vụ tại công ty.",
            "Đọc mô tả ngắn trên từng Card rồi tiếp tục; lựa chọn sẽ đi cùng Booking để Admin, Quản lý và KTV chuẩn bị phù hợp.",
        ],
        "Màn hình chuyển đúng sang luồng tương ứng, không mất trạng thái khi bấm quay lại trong cùng hành trình.",
        "02-khach-hang-dat-lich-nhanh.png",
        "Hình 2. Popup Hình thức đặt lịch — bốn lựa chọn được phân cấp bằng màu sắc.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "3. Chọn dịch vụ, cơ sở, KTV, ngày và giờ",
        "Trình tự cơ sở → KTV → ngày → giờ giúp hiển thị đúng năng lực phục vụ và tránh xung đột KTV/giường.",
        [
            "Chọn dịch vụ và số người; hệ thống lấy đúng thời lượng của từng dịch vụ.",
            "Chọn Cơ sở 1 hoặc Cơ sở 2, sau đó chọn KTV cụ thể hoặc Ngẫu nhiên.",
            "Quan sát bảng khả dụng: xanh là còn nhận được, đỏ là bận/hết chỗ; chỉ chọn ô xanh.",
            "Chọn ngày và giờ. Một ca 90 phút lúc 10:00 giữ đồng thời KTV và giường đến 11:30.",
            "Kiểm tra thông tin khách; nếu đã đăng nhập, dữ liệu tài khoản được tự điền nhưng vẫn có thể chỉnh sửa.",
        ],
        "Tổng tiền tạm tính, thời lượng, cơ sở, KTV và thời điểm được ghi rõ trước bước đặt cọc.",
        "03-khach-hang-chon-dich-vu-co-so.png",
        "Hình 3. Luồng đặt lịch — lựa chọn dịch vụ và cơ sở trước khi chốt nguồn lực.",
        "Cơ sở mở 09:00–24:00. Khung 23:00 chỉ nhận dịch vụ 60 phút; dịch vụ 90 phút không được bắt đầu lúc 23:00.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "4. Ưu đãi, đặt cọc và xác nhận chỗ",
        "Khoản cọc giữ chỗ cho khách nhưng vẫn tách minh bạch phần thanh toán tại cơ sở. Voucher chỉ giảm phần còn lại, không làm thay đổi tiền cọc nền tảng.",
        [
            "Chọn voucher đủ điều kiện; voucher chào mừng chỉ dùng một lần trên tài khoản hợp lệ.",
            "Đọc Hóa đơn tạm tính và đánh dấu đồng ý Điều khoản, Chính sách dữ liệu, Chính sách đặt lịch/đặt cọc.",
            "Thanh toán khoản cọc bằng QR ngân hàng của nền tảng.",
            "Sau đối soát, hệ thống xác nhận tự động trong 1–2 giây nếu AI tự động xác nhận được bật; nếu không, Booking chuyển sang hàng chờ Admin/Quản lý duyệt.",
        ],
        "Đơn chuyển sang Đã đặt cọc/Đã xác nhận, có nút xem Đơn và sẵn sàng dùng QR tại cơ sở.",
        "04-khach-hang-coc-va-uu-dai.png",
        "Hình 4. Hóa đơn tạm tính — tách Tổng cộng, Cọc nền tảng và Còn lại tại cơ sở.",
        "Cọc = 10% giá trị Bill ban đầu trước ưu đãi. Còn lại = 90% giá trị Bill ban đầu − ưu đãi. Ví dụ giá gốc 250.000đ, ưu đãi 50.000đ: cọc 25.000đ; còn lại 175.000đ.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "5. Đơn của tôi và Thu–Chi cá nhân",
        "Một điểm tra cứu thống nhất cho lịch đã đặt, khoản cọc, phần còn lại, lịch sử sử dụng và thu nhập Affiliate.",
        [
            "Bấm biểu tượng $ trên Topbar hoặc CTA Đơn của tôi.",
            "Chọn Tổng quan để xem dòng tiền; chọn Chi tiết để mở từng Bill; dùng bộ lọc ngày/tuần/tháng.",
            "Từ một đơn đủ điều kiện, chọn Quét QR sử dụng dịch vụ hoặc mở đơn sau khi đã quét QR cơ sở.",
            "Nếu đổi thiết bị, đăng nhập đúng số điện thoại để khôi phục lịch đã đặt.",
        ],
        "Khoản cọc, phần còn lại và trạng thái dịch vụ không bị trùng; dữ liệu Affiliate được trình bày riêng khỏi Bill dịch vụ.",
        "06-khach-hang-thu-chi.png",
        "Hình 5. Thu–Chi của tôi — tổng hợp giao dịch cá nhân và lối tắt tới Đơn của tôi.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "6. QR check-in, đồng hồ phục vụ và checkout",
        "QR nối đúng khách–Bill–cơ sở–KTV. Đồng hồ phục vụ giúp khách, KTV và quản lý cùng quan sát một trạng thái thời gian thực.",
        [
            "Tại cơ sở, bấm biểu tượng QR trên Topbar và cho phép Camera.",
            "Quét QR dán tại Cơ sở 1/Cơ sở 2 hoặc QR riêng của KTV. Hệ thống chỉ mở các Bill/thẻ đủ điều kiện của tài khoản hiện tại.",
            "Chọn Sử dụng dịch vụ để check-in; đồng hồ bắt đầu theo đúng thời lượng của dịch vụ.",
            "Trong quá trình chạy, có thể mở lại đồng hồ từ banner Đang phục vụ hoặc Card Thu–Chi.",
            "Nếu cần về sớm, chọn Checkout sớm và xác nhận. Khi hết ca, thanh toán phần còn lại qua QR ngân hàng của cơ sở, rồi đánh giá dịch vụ.",
        ],
        "KTV/giường chuyển sang bận; Admin/Quản lý nhìn thấy ca đang phục vụ và thời gian còn lại; sau checkout trạng thái trở về rảnh.",
        "11-khach-hang-ho-so-ktv.png",
        "Hình 6. Hồ sơ KTV — thông tin năng lực giúp khách lựa chọn trước khi đặt lịch.",
        "Tip hoàn toàn tùy tâm, khách trao trực tiếp cho KTV; không cộng vào Bill, QR thanh toán dịch vụ hay doanh thu cơ sở. Trong đồng hồ chỉ nhắc khéo, không đặt mức sàn.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "7. Trung tâm thông báo khách hàng",
        "Thông báo tập trung giúp khách biết ngay lịch được xác nhận/từ chối, khoản cọc đã nhận, trạng thái dịch vụ, lời mời và khuyến mãi.",
        [
            "Bấm Chuông trên Topbar.",
            "Lọc theo Tất cả, Chưa đọc, Cơ sở, lời mời/quan hệ và khuyến mãi.",
            "Bấm một thông báo để mở đúng đơn hoặc hành động liên quan; dùng Đọc tất cả khi đã xử lý xong.",
        ],
        "Số chưa đọc giảm đúng; nội dung chỉ hiển thị thông tin phù hợp vai trò khách, không lộ ghi chú nội bộ.",
        "08-khach-hang-thong-bao.png",
        "Hình 7. Trung tâm thông báo khách hàng — lọc và đánh dấu đã đọc theo ngữ cảnh.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "8. Affiliate — giới thiệu và theo dõi thu nhập",
        "Mỗi người có mã/QR riêng; hệ thống gắn nguồn giới thiệu vào đơn phát sinh để đo chuyển đổi và tính thưởng minh bạch.",
        [
            "Mở Affiliate trên thanh điều hướng.",
            "Sao chép mã, chia sẻ đường link hoặc gửi QR cho bạn bè/đối tác.",
            "Người được giới thiệu mở link/QR, đặt lịch và hoàn tất điều kiện của chương trình.",
            "Theo dõi số đã mời, số thành công, tỷ lệ chốt, hạng và thu nhập theo tháng.",
        ],
        "Mã nguồn được giữ xuyên hành trình; thu nhập Affiliate hiển thị riêng và có thể đối chiếu theo giao dịch đủ điều kiện.",
        "07-affiliate-trung-tam.png",
        "Hình 8. Trung tâm Affiliate — mã giới thiệu, QR, tiến độ hạng và báo cáo thu nhập.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "9. Tuệ Tâm Business — dịch vụ tại doanh nghiệp",
        "Số hóa từ yêu cầu, báo giá, cọc, điều phối đoàn KTV đến thời gian phục vụ và thanh toán công nợ.",
        [
            "Chọn Business trong Hình thức đặt lịch; điền tên công ty, mã số thuế, địa điểm, đầu mối, quy mô và nhu cầu.",
            "Hệ thống lập báo giá gồm dịch vụ và phí di chuyển; khách thanh toán cọc nền tảng để xác nhận.",
            "Admin/Quản lý duyệt hoặc AI tự xác nhận, phân công KTV Business trưởng và đoàn KTV.",
            "Tại doanh nghiệp, đầu mối quét QR của KTV Business trưởng để bắt đầu phiên; chuông nhắc khi gần kết thúc.",
            "Kết thúc bằng QR của KTV trưởng, đối soát thời gian/địa điểm, thanh toán phần còn lại cho cơ sở và gửi đánh giá.",
        ],
        "Booking Business xuất hiện trong khu vực Booking và lịch đoàn; Admin/Quản lý/KTV trưởng cùng thấy công nợ, thời gian, địa điểm, nhân sự và đánh giá.",
        "09-business-dat-dich-vu.png",
        "Hình 9. Đặt Tuệ Tâm Business — tiếp nhận đầy đủ thông tin doanh nghiệp và địa điểm triển khai.",
        "Cọc Business cũng theo quy tắc 10% giá trị báo giá ban đầu trước ưu đãi; phần còn lại bằng giá trị sau ưu đãi trừ tiền cọc.",
    )
    page_break(doc)

    doc.add_heading("PHẦN III — KỸ THUẬT VIÊN", level=1)
    add_feature_page(
        doc,
        "1. Lịch làm việc cá nhân",
        "KTV nhìn theo tháng/tuần, biết mỗi ngày có bao nhiêu lịch tại cơ sở và bao nhiêu đoàn Business; giảm bỏ sót ca.",
        [
            "Đăng nhập tài khoản KTV được cấp; trang mặc định là Lịch.",
            "Dùng mũi tên đổi tháng hoặc bấm Hôm nay.",
            "Mỗi ô ngày có hai con số: lịch tại cơ sở và đoàn Business; bấm ngày để xem chi tiết bên dưới.",
            "Mở hồ sơ ca để xem khách, dịch vụ, yêu cầu riêng, giường/phòng và căn dặn của IQ Care.",
        ],
        "KTV biết rõ thời gian, nơi làm, loại khách và việc cần chuẩn bị trước ca.",
        "23-ktv-lich-ca-nhan.png",
        "Hình 10. Lịch cá nhân KTV — hai lớp dữ liệu cơ sở và Business trên cùng lịch tháng.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "2. Điều phối tại cơ sở",
        "Màn hình tức thời cho biết ai rảnh, ai bận, ca nào sắp bắt đầu và còn bao lâu để phân công hợp lý.",
        [
            "Mở tab Điều phối.",
            "Quan sát Card KTV: xanh là rảnh; đỏ là bận và kèm thời gian còn lại/ngữ cảnh ca.",
            "Đọc Dòng điều phối hôm nay để biết khách, dịch vụ, giờ, KTV và giường đã xếp.",
            "Nếu có thay đổi, trao đổi với Lễ tân/Quản lý; KTV không tự ý sửa dữ liệu của đồng nghiệp.",
        ],
        "Lịch, phòng/giường và trạng thái KTV nhất quán giữa KTV, Lễ tân, Quản lý và Admin.",
        "24-ktv-dieu-phoi.png",
        "Hình 11. Điều phối KTV — trạng thái bận/rảnh và lịch trong ngày theo cơ sở.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "3. Thu nhập và lịch sử Tip",
        "Tách thu nhập trách nhiệm khỏi dữ liệu Tip lịch sử; KTV dễ đối chiếu mà không làm sai doanh thu dịch vụ.",
        [
            "Mở tab Thu nhập.",
            "Xem thu nhập trách nhiệm và các khoản đã được ghi nhận theo kỳ.",
            "Dữ liệu Tip chỉ giữ để đối chiếu lịch sử nếu có; Tip mới do khách trao trực tiếp và không đi qua nền tảng.",
            "Nếu lệch số, gửi yêu cầu cho Quản lý cơ sở kèm ngày, khách và ca phục vụ.",
        ],
        "KTV thấy rõ khoản nào thuộc thu nhập, khoản nào chỉ là lịch sử Tip; không trộn vào Bill dịch vụ.",
        "25-ktv-thu-nhap.png",
        "Hình 12. Thu nhập KTV — minh bạch khoản cá nhân và lịch sử Tip.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "4. Hồ sơ năng lực và quy trình duyệt",
        "Hồ sơ KTV tạo niềm tin cho khách và giúp hệ thống gợi ý đúng thế mạnh, nhưng chỉ công khai sau khi Admin duyệt.",
        [
            "Mở tab Tôi, tải ảnh đại diện rõ mặt và phù hợp quy chuẩn.",
            "Viết giới thiệu ngắn, mô tả phong cách phục vụ; nhập tối đa sáu điểm mạnh, mỗi dòng một điểm.",
            "Bấm Gửi nội dung cập nhật để duyệt.",
            "Admin kiểm tra ảnh, ngôn từ, chứng chỉ và thế mạnh; sau duyệt hồ sơ mới xuất hiện cho khách.",
        ],
        "Trạng thái hồ sơ hiển thị rõ Chờ duyệt/Đã công khai/Yêu cầu chỉnh sửa.",
        "26-ktv-ho-so-ca-nhan.png",
        "Hình 13. Hồ sơ KTV — ảnh, giới thiệu, thế mạnh và CTA gửi duyệt.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "5. QR riêng của KTV",
        "QR KTV nối khách vào đúng Bill đã phân công và tự cập nhật KTV đang bận, thời gian còn lại cho quản lý.",
        [
            "Bấm QR trên Topbar KTV.",
            "Đưa mã cho khách quét bằng Camera trong Tuệ Tâm Care.",
            "Khách chọn Bill đủ điều kiện và bắt đầu dịch vụ; QR không mở Bill của khách khác.",
            "Khi đổi nhân sự hoặc nghi ngờ lộ mã, Admin/Quản lý cấp lại QR; mã cũ hết hiệu lực ngay.",
        ],
        "Ca được gắn đúng KTV; trạng thái bận/rảnh đồng bộ; khách mở được đồng hồ phục vụ.",
        "27-ktv-qr-ca-nhan.png",
        "Hình 14. QR riêng KTV — mở từ Topbar, tải xuống và sử dụng tại điểm phục vụ.",
    )
    page_break(doc)

    doc.add_heading("PHẦN IV — LỄ TÂN, QUẢN LÝ & ADMIN", level=1)
    add_feature_page(
        doc,
        "1. Tổng quan vận hành theo phạm vi",
        "Một màn hình cân bằng giữa Booking, khách, công suất, KTV và tài chính. Bộ lọc điều khiển đồng bộ toàn bộ Card.",
        [
            "Chọn Hôm nay/Tuần/Tháng/Năm/Từ ngày/Bất kỳ.",
            "Admin chọn Toàn hệ thống hoặc cơ sở; Quản lý/Lễ tân bị khóa đúng cơ sở được giao.",
            "Chọn ca Cả ngày/Sáng/Chiều/Tối để xem lịch trong phạm vi.",
            "Bấm Card để đi sâu tới Booking, khách, công suất hoặc tài chính tương ứng.",
        ],
        "Các con số thay đổi đồng bộ theo cùng bộ lọc; Quản lý cơ sở không nhìn thấy dữ liệu ngoài phạm vi.",
        "33-quan-ly-co-so-tong-quan.png",
        "Hình 15. Tổng quan Quản lý Cơ sở 1 — KPI vận hành và tài chính trong phạm vi được cấp.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "2. Booking, xác nhận và AI tự động điều phối",
        "Booking là hàng chờ yêu cầu; Lịch chỉ chứa ca đã xác nhận/xếp nguồn lực. Hai màn hình tách rõ để tránh nhầm trạng thái.",
        [
            "Mở Booking và chọn ngày có số lượng yêu cầu.",
            "Xem loại Cá nhân/Mời bạn/Mời sếp/Business, khoản cọc và yêu cầu riêng.",
            "Nếu AI tự động xác nhận bật, hệ thống kiểm tra KTV/giường rồi xác nhận trong 1–2 giây; nếu tắt, chọn Xác nhận hoặc Từ chối kèm lý do.",
            "Sau xác nhận, kiểm tra Lịch để bảo đảm ca đã có KTV, giường/phòng và không xung đột thời lượng.",
        ],
        "Booking mới nằm trên đầu; khách nhận thông báo; ca xác nhận xuất hiện trong Lịch và sơ đồ nguồn lực.",
        "14-admin-lich.png",
        "Hình 16. Lịch Admin — lịch tháng và chi tiết ca đã được xác nhận/xếp lịch.",
        "Một Booking 90 phút giữ nguồn lực toàn bộ 90 phút. Hệ thống không cho đặt chồng KTV hoặc giường trong khoảng đó.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "3. CRM khách hàng tại quầy",
        "Lễ tân và quản lý nhìn được lịch sử ghé, chi tiêu, KTV quen và phân nhóm để chăm sóc cá nhân hóa.",
        [
            "Mở IQ Care → Khách hàng hoặc dùng Tạo khách nhanh khi khách đến trực tiếp.",
            "Tìm bằng tên/số điện thoại; lọc VIP, thân thiết, thẻ dài hạn, doanh nghiệp, Affiliate…",
            "Chọn Tóm tắt để quét nhanh; chọn Đầy đủ để xem timeline, đơn, phản hồi và mối quan hệ.",
            "Không tạo hồ sơ trùng số điện thoại; cập nhật ghi chú phục vụ ở đúng trường nghiệp vụ.",
        ],
        "Khách được nhận diện nhất quán ở Booking, Lịch, Bill và chăm sóc lại; ghi chú nội bộ không hiển thị cho khách.",
        "34-le-tan-crm-khach-hang.png",
        "Hình 17. CRM tại quầy — phân nhóm khách, lượt ghé, chi tiêu và KTV quen.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "4. Ghi nhận chi phí phát sinh",
        "Quản lý/Admin nhập mọi khoản chi có chứng từ; AI hỗ trợ phân loại để Thu–Chi và lãi/lỗ cập nhật đúng cơ sở/hệ thống.",
        [
            "Bấm dấu + trên Topbar quản trị. Chức năng này chỉ có ở Quản lý cơ sở và Admin/Chủ.",
            "Nhập Số tiền trước, nội dung công việc, danh mục và phạm vi hạch toán: Cơ sở 1, Cơ sở 2 hoặc Chi cho hệ thống.",
            "Chọn ngày, nhà cung cấp/người nhận; tải ảnh Bill để AI nhận diện và đối soát.",
            "Kiểm tra danh mục gợi ý như cơ sở vật chất, khấu hao, lương, thưởng, mặt bằng, điện nước, marketing…",
            "Bấm Ghi nhận đã chi. Không ghi nhận khi thiếu căn cứ hoặc phạm vi chưa đúng.",
        ],
        "Khoản chi xuất hiện trong Chi phí, đúng cơ sở và kỳ báo cáo; nhật ký lưu người tạo và thời điểm.",
        "18-admin-ghi-nhan-chi-phi.png",
        "Hình 18. Ghi nhận chi phí phát sinh — số tiền, nội dung, danh mục, phạm vi và ảnh Bill AI.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "5. Trung tâm tài chính",
        "Tập trung Doanh thu, Chi phí, Lãi tạm tính và dữ liệu Tip lịch sử; cho phép đi từ tổng quan đến từng Bill/khoản chi.",
        [
            "Bấm biểu tượng $ trên Topbar.",
            "Chọn kỳ, cơ sở và khung giờ; dùng tìm kiếm theo khách, số điện thoại, dịch vụ hoặc khoản chi.",
            "Bấm Card Doanh thu/Chi phí/Lãi tạm tính để lọc danh sách chi tiết.",
            "Bấm Cơ sở 1/Cơ sở 2 để xem cấu thành Thu–Chi riêng; mở từng Bill để xem Tổng, Đã cọc, Còn lại/Đã thanh toán.",
            "Đối chiếu Tip ở khu vực riêng chỉ cho dữ liệu lịch sử; không cộng vào Bill dịch vụ mới.",
        ],
        "Tổng thu − Tổng chi = Lãi tạm tính theo cùng phạm vi; số liệu có thể truy nguyên đến giao dịch gốc.",
        "16-admin-trung-tam-tai-chinh.png",
        "Hình 19. Trung tâm tài chính — bộ lọc, KPI và lãi/lỗ theo cơ sở.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "6. Quản lý QR hệ thống",
        "Quản lý tập trung QR cơ sở, QR KTV và QR Business; cấp lại có kiểm soát để tránh dùng mã cũ.",
        [
            "Mở Trung tâm quản lý QR từ Topbar/IQ Care.",
            "Chọn tab Cơ sở, KTV hoặc Business; tìm theo tên.",
            "Tải QR để in/dán, sao chép Link hoặc Gửi cho đầu mối triển khai.",
            "Khi cấp lại, xác nhận phạm vi; hệ thống tăng phiên và vô hiệu ngay mã trước đó.",
        ],
        "Mỗi mã có đúng chủ thể và phiên; lịch sử cấp lại giúp Admin kiểm soát rủi ro.",
        "19-admin-quan-ly-qr.png",
        "Hình 20. Trung tâm quản lý QR — QR cơ sở, KTV, Business và thao tác cấp lại.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "7. Thông báo quản trị",
        "Admin/Quản lý nhận thông báo theo cơ sở và nghiệp vụ: Booking, thanh toán, vận hành, CRM VIP, Business và chi hệ thống.",
        [
            "Bấm Chuông; chọn Tất cả hoặc Chưa đọc.",
            "Lọc theo nghiệp vụ và phạm vi Toàn hệ thống/Cơ sở 1/Cơ sở 2/Chi hệ thống.",
            "Mở thông báo để đến đúng hồ sơ cần xử lý; đánh dấu đã đọc sau khi hoàn tất.",
            "Dùng Đọc tất cả chỉ khi đã rà soát toàn bộ cảnh báo quan trọng.",
        ],
        "Số chưa đọc, phạm vi và đường dẫn hành động nhất quán; nội dung phù hợp vai trò quản trị.",
        "17-admin-thong-bao.png",
        "Hình 21. Thông báo quản trị — bộ lọc nghiệp vụ, cơ sở và trạng thái đọc.",
        "Một số thông báo Tip cũ có thể còn trong dữ liệu UAT để đối chiếu lịch sử; giao dịch mới không đưa Tip vào Bill.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "8. IQ Care và cấu hình hệ thống",
        "IQ Care gom các module theo nghiệp vụ; Trung tâm cấu hình lưu CSDL, phân phạm vi và ghi nhật ký mỗi lần thêm/sửa/xóa.",
        [
            "Mở IQ Care để vào Tổng quan, Đặt lịch, Công suất, Phòng & giường, Khách hàng, KTV, Voucher, Business, Báo cáo…",
            "Admin/Chủ mở Cấu hình, tìm theo từ khóa hoặc nhóm nghiệp vụ.",
            "Bật/tắt AI tự động xác nhận, trợ lý chat và thông báo; sửa giờ hoạt động, tỷ lệ cọc, MFA và chính sách vận hành.",
            "Mọi thay đổi phải có lý do, phạm vi và người chịu trách nhiệm; kiểm tra tác động trên UAT trước khi áp dụng production.",
        ],
        "Giá trị cấu hình được đọc bởi luồng thật; thao tác có nhật ký và chỉ người đúng quyền mới sửa được.",
        "22-admin-cau-hinh-he-thong.png",
        "Hình 22. Trung tâm cấu hình — tìm kiếm, nhóm nghiệp vụ và thao tác thêm/sửa/xóa.",
    )
    page_break(doc)

    doc.add_heading("PHẦN V — NHÀ ĐẦU TƯ", level=1)
    add_feature_page(
        doc,
        "1. Tổng quan danh mục đang vận hành",
        "Nhà đầu tư xem vốn, doanh thu, chi phí, lợi nhuận, phần lợi nhuận của mình và dự kiến hoàn vốn trong phạm vi đã đầu tư.",
        [
            "Đăng nhập tài khoản Nhà đầu tư; giao diện mặc định ở Tổng quan và chỉ xem.",
            "Chọn Hôm nay/Tuần/Tháng/Quý/Năm/Tùy chỉnh.",
            "Chọn Tất cả cơ sở hoặc từng cơ sở đã đầu tư độc lập.",
            "Bấm Card cơ sở để xem chi tiết; đọc dải dự kiến hoàn vốn AI và độ tin cậy.",
        ],
        "Cơ sở đang vận hành được tách rõ khỏi cơ hội mới; dữ liệu ngoài phạm vi đầu tư không hiển thị.",
        "28-nha-dau-tu-tong-quan.png",
        "Hình 23. Trung tâm Nhà đầu tư — danh mục hiện hữu, bộ lọc thời gian và phạm vi vốn.",
        "Dự kiến hoàn vốn là ước tính theo tiến độ hiện tại, hiển thị dải ±10%; không phải cam kết lợi nhuận.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "2. Hiệu quả tài chính",
        "Biểu đồ so sánh Doanh thu–Chi phí–Lợi nhuận và Card cơ sở giúp nhà đầu tư đi từ xu hướng đến cấu thành chi tiết.",
        [
            "Mở tab Hiệu quả; giữ cùng kỳ và phạm vi cần so sánh.",
            "Bấm từng cột ngày/tháng để mở số liệu cụ thể.",
            "Bấm Card Cơ sở 1/Cơ sở 2 để xem doanh thu thuần, chi phí, lợi nhuận và biên lợi nhuận.",
            "Đối chiếu với bản tin điều hành nếu có biến động lớn.",
        ],
        "Biểu đồ có biến thiên thực tế, không san đều; tổng Card khớp với kỳ và cơ sở đã chọn.",
        "29-nha-dau-tu-hieu-qua.png",
        "Hình 24. Hiệu quả tài chính — biểu đồ và bộ lọc độc lập theo cơ sở đã đầu tư.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "3. Cơ hội đầu tư mới",
        "Pipeline chỉ chứa địa điểm đang khảo sát/thẩm định; chưa giải ngân, chưa có quyền sở hữu và không cộng vào hiệu quả cơ sở hiện hữu.",
        [
            "Mở tab Cơ hội mới.",
            "Đọc nhu cầu vốn, mức tham gia tối thiểu, thời gian hồi vốn dự kiến, tiến độ thẩm định và tài liệu công bố.",
            "Chọn Đăng ký quan tâm nếu muốn Investor Care liên hệ; hành động này chưa phải cam kết góp vốn.",
            "Theo dõi Chuông để nhận hồ sơ mới, cập nhật thầu lại/khảo sát và quyết định của Admin.",
        ],
        "Cơ hội mới luôn có nhãn pipeline và trạng thái thẩm định; không bị nhầm với Cơ sở 1/Cơ sở 2 đang vận hành.",
        "30-nha-dau-tu-co-hoi-moi.png",
        "Hình 25. Cơ hội đầu tư mới — pipeline tách biệt hoàn toàn với tài sản đang hoạt động.",
    )
    page_break(doc)

    add_feature_page(
        doc,
        "4. Đặc quyền và bản tin Nhà đầu tư",
        "Ngoài dữ liệu tài chính, nhà đầu tư được quản lý quyền lợi sức khỏe, tiếp đón đối tác và tiếp cận hồ sơ sớm theo chính sách riêng.",
        [
            "Mở tab Đặc quyền để xem quyền lợi đang kích hoạt và đầu mối Investor Care.",
            "Bấm Chuông để lọc Cơ hội mới, Đặc quyền và Điều hành.",
            "Mở từng bản tin; đánh dấu đã đọc sau khi xem nội dung và tài liệu liên quan.",
            "Liên hệ Investor Care nếu cần đặt lịch chăm sóc đặc quyền hoặc tiếp đón đối tác.",
        ],
        "Quyền lợi phi tài chính tách khỏi lợi nhuận; thông tin tài chính riêng tư không dùng trong luồng tiếp đón.",
        "31-nha-dau-tu-dac-quyen.png",
        "Hình 26. Đặc quyền Nhà đầu tư — quyền lợi sức khỏe, đối tác và tiếp cận cơ hội sớm.",
    )
    page_break(doc)

    doc.add_heading("PHẦN VI — LOGIC LIÊN VAI TRÒ", level=1)
    doc.add_heading("1. Chuỗi trạng thái Booking tiêu chuẩn", level=2)
    add_flow_table(doc, [
        ("Gửi yêu cầu", "Xem tạm tính và đặt cọc", "Nhận Booking mới", "Ghi nhận yêu cầu; chưa ghi doanh thu hoàn tất"),
        ("Cọc đối soát", "Nhận xác nhận giữ chỗ", "AI/QL xếp KTV và giường", "Ghi cọc nền tảng; tạo thông báo"),
        ("Check-in QR", "Chọn Bill/thẻ để sử dụng", "KTV/giường chuyển bận", "Mở phiên dịch vụ và nhật ký thời gian"),
        ("Đang phục vụ", "Xem đồng hồ/checkout sớm", "Theo dõi ca và thời gian còn lại", "Hiển thị tức thời ở vận hành"),
        ("Checkout", "Xác nhận kết thúc", "Nguồn lực trở về rảnh", "Chốt thời gian; mở phần thanh toán còn lại"),
        ("Thanh toán nốt", "Chuyển QR cho cơ sở", "Đối soát và kết thúc Bill", "Ghi doanh thu dịch vụ; không gồm Tip"),
        ("Đánh giá", "Chấm điểm/góp ý", "KTV nhận phản hồi phù hợp", "CRM, chất lượng và báo cáo được cập nhật"),
    ])
    doc.add_heading("2. Công thức tài chính bắt buộc", level=2)
    add_callout(doc, "Công thức chuẩn", "Cọc nền tảng = 10% giá trị Bill ban đầu trước ưu đãi. Phần còn lại tại cơ sở = 90% giá trị Bill ban đầu − tổng ưu đãi hợp lệ. Tip = tùy tâm, khách trao trực tiếp cho KTV, ngoài Bill và ngoài QR thanh toán dịch vụ.", PALE_GOLD, "A66A08")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2200, 2100, 2100, 2960])
    for idx, txt in enumerate(("Giá gốc", "Ưu đãi", "Cọc nền tảng", "Còn lại tại cơ sở")):
        c = table.rows[0].cells[idx]
        c.text = txt
        shade(c, DARK_BROWN)
        cell_border(c, DARK_BROWN)
        set_cell_margins(c)
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
    set_repeat_table_header(table.rows[0])
    for data in (("250.000đ", "0đ", "25.000đ", "225.000đ"), ("250.000đ", "50.000đ", "25.000đ", "175.000đ"), ("1.150.000đ", "100.000đ", "115.000đ", "935.000đ")):
        cells = table.add_row().cells
        for idx, text in enumerate(data):
            cells[idx].text = text
            shade(cells[idx], CREAM)
            cell_border(cells[idx])
            set_cell_margins(cells[idx])
    doc.add_heading("3. Quy tắc lịch–KTV–giường", level=2)
    add_bullets(doc, [
        "Khoảng bận bắt đầu tại giờ hẹn và kết thúc sau đúng thời lượng dịch vụ; cùng khoảng thời gian được giữ cho cả KTV và giường/phòng.",
        "Không cho phép hai Booking giao nhau trên cùng một KTV hoặc cùng một giường, kể cả khi khác khách hoặc khác kênh đặt.",
        "Mở cửa 09:00, đóng cửa 24:00; ca nhận khách cuối 23:00 và chỉ nhận dịch vụ 60 phút.",
        "Cơ sở 1: 28 giường, 8 KTV. Cơ sở 2: 32 giường, 8 KTV. Admin cập nhật nhân sự và lịch làm trước khi chạy thật.",
        "KTV Business đang đi đoàn phải hiển thị bận trong điều phối cơ sở để tránh xếp trùng.",
    ])
    doc.add_heading("4. Chính sách đổi lịch/không đến", level=2)
    add_bullets(doc, [
        "Mỗi khách được đổi lịch miễn phí một lần trong tháng.",
        "Từ lần đổi thứ hai trong cùng tháng, áp dụng chi phí cọc của lần trước theo chính sách hiện hành.",
        "Khách không đến được nhắn lại tự động tối đa một lần trong tháng và được nhắc lịch khéo léo cho lần tiếp theo.",
        "Nếu khách đến và sử dụng đúng lịch thì không bị phạt cọc.",
    ])
    page_break(doc)

    doc.add_heading("PHẦN VII — CHECKLIST VẬN HÀNH THỬ", level=1)
    doc.add_heading("1. Trước giờ mở cửa", level=2)
    add_bullets(doc, [
        "Xác nhận 60 giường/phòng và 16 KTV đã được cấu hình đúng trạng thái, ca làm và cơ sở.",
        "Kiểm tra QR Cơ sở 1, Cơ sở 2 và QR KTV quét được trên ít nhất hai điện thoại.",
        "Kiểm tra tài khoản ngân hàng nền tảng cho cọc và tài khoản cơ sở cho phần còn lại; không dùng tài khoản demo khi chạy thật.",
        "Kiểm tra AI tự động xác nhận, thông báo và tác vụ nền; bảo đảm không có cảnh báo heartbeat.",
        "Bật MFA cho Admin/Chủ và Quản lý; rà soát người dùng đã nghỉ việc, QR đã cấp và quyền theo cơ sở.",
    ])
    doc.add_heading("2. Kịch bản test xuyên suốt", level=2)
    add_steps(doc, [
        "Tạo một khách mới, nhận voucher chào mừng và đặt Foot Massage 60 phút.",
        "Thanh toán cọc; kiểm tra Booking mới lên đầu và khách nhận thông báo xác nhận.",
        "Kiểm tra Lịch, KTV và giường được giữ đúng 60 phút.",
        "Quét QR cơ sở, chọn Bill, bắt đầu đồng hồ; kiểm tra KTV/giường chuyển bận ở Quản lý/Admin.",
        "Checkout sớm hoặc hết giờ; thanh toán phần còn lại; kiểm tra Bill hoàn tất và nguồn lực trở về rảnh.",
        "Thực hiện một giao dịch Affiliate và một Booking Business; kiểm tra thông báo ở tất cả vai trò liên quan.",
        "Ghi một khoản chi có ảnh Bill; đối chiếu Tổng chi và Lãi tạm tính đúng cơ sở.",
        "Đăng nhập Nhà đầu tư và xác nhận dữ liệu chỉ xem, đúng phạm vi, tách cơ hội mới khỏi cơ sở vận hành.",
    ])
    doc.add_heading("3. Dấu hiệu cần dừng giao dịch và báo quản trị", level=2)
    add_bullets(doc, [
        "Cọc đã trừ tiền nhưng sau thời gian đối soát vẫn không có Booking/Thông báo.",
        "Khách quét QR thấy Bill của người khác hoặc ghi chú nội bộ.",
        "Một KTV/giường bị xếp chồng ca; dịch vụ vượt quá 24:00 vẫn được nhận.",
        "Tip xuất hiện trong Bill, QR thanh toán dịch vụ hoặc doanh thu cơ sở.",
        "Quản lý nhìn thấy cơ sở không được giao; Lễ tân ghi được chi phí; Nhà đầu tư sửa được dữ liệu.",
        "Tổng thu, tổng chi và lãi tạm tính không truy nguyên được đến giao dịch cụ thể.",
    ])
    add_callout(doc, "Khi báo lỗi", "Gửi vai trò đang dùng, thời điểm, số điện thoại/mã đơn đã che bớt, các bước vừa thao tác, ảnh màn hình và kết quả mong đợi. Không gửi mật khẩu, token, mã khôi phục MFA hoặc ảnh chứng từ chứa dữ liệu nhạy cảm qua kênh công khai.", PALE_RED, BURGUNDY)
    page_break(doc)

    doc.add_heading("PHỤ LỤC — TRA CỨU NHANH", level=1)
    doc.add_heading("Các điểm vào chính", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2300, 3300, 3760])
    for idx, text in enumerate(("Biểu tượng/Khu vực", "Dùng cho", "Kết quả")):
        c = table.rows[0].cells[idx]
        c.text = text
        shade(c, TEAL)
        cell_border(c, TEAL)
        set_cell_margins(c)
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
    set_repeat_table_header(table.rows[0])
    entries = [
        ("+ khách hàng", "Chọn hình thức đặt lịch", "Cá nhân/Mời bạn/Mời sếp/Business"),
        ("QR Topbar", "Quét QR cơ sở/KTV/Business", "Mở Bill đủ điều kiện và check-in"),
        ("$ khách hàng", "Thu–Chi và Đơn của tôi", "Tra cứu cọc, còn lại, trạng thái"),
        ("Chuông", "Thông báo theo vai trò", "Lọc, mở hành động, đánh dấu đọc"),
        ("+ quản trị", "Ghi nhận khoản chi", "Hạch toán đúng cơ sở/hệ thống"),
        ("$ quản trị", "Trung tâm tài chính", "Doanh thu, chi phí, lãi/lỗ, Bill"),
        ("IQ Care", "Module quản trị", "Khách, KTV, dịch vụ, cấu hình, báo cáo"),
    ]
    for i, entry in enumerate(entries):
        cells = table.add_row().cells
        for j, text in enumerate(entry):
            cells[j].text = text
            shade(cells[j], "FFFFFF" if i % 2 == 0 else CREAM)
            cell_border(cells[j])
            set_cell_margins(cells[j])

    doc.add_heading("Thông điệp cốt lõi", level=2)
    add_callout(
        doc,
        "Một nền tảng — một chuỗi dữ liệu",
        "Mỗi thao tác của khách phải tạo ra trạng thái có ý nghĩa cho KTV, Lễ tân, Quản lý, Admin và tài chính; mỗi vai trò chỉ nhìn thấy phần dữ liệu cần thiết để phục vụ tốt, vận hành đúng và ra quyết định minh bạch.",
        PALE_GREEN,
        GREEN,
    )
    p = doc.add_paragraph("TUỆ TÂM CARE • TẬN TÂM TRONG TRẢI NGHIỆM • MINH BẠCH TRONG VẬN HÀNH", style="Guide Eyebrow")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
