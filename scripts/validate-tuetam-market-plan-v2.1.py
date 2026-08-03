from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from zipfile import ZipFile

from docx import Document


DOCX = Path(
    r"C:\Users\DELL\Desktop\Kyoto Masage\artifacts\Tue_Tam_Care_Nghien_Cuu_Thi_Truong_Ke_Hoach_Kinh_Doanh_Tinh_Kha_Thi_v2.1.docx"
)
GMV = Decimal("52200000000")
SHARES = {
    "KTV trực tiếp": Decimal("0.60"),
    "Trưởng đoàn": Decimal("0.05"),
    "Nền tảng": Decimal("0.20"),
    "Giám đốc phân phối cấp Quận": Decimal("0.05"),
    "Affiliate trực tiếp": Decimal("0.10"),
}


def main() -> None:
    assert DOCX.exists(), DOCX
    assert sum(SHARES.values()) == Decimal("1.00")

    expected = {
        role: GMV * share
        for role, share in SHARES.items()
    }
    assert expected["KTV trực tiếp"] == Decimal("31320000000")
    assert expected["Trưởng đoàn"] == Decimal("2610000000")
    assert expected["Nền tảng"] == Decimal("10440000000")
    assert expected["Giám đốc phân phối cấp Quận"] == Decimal("2610000000")
    assert expected["Affiliate trực tiếp"] == Decimal("5220000000")
    assert sum(expected.values()) == GMV

    document = Document(DOCX)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    required = [
        "KTV trực tiếp thực hiện",
        "Trưởng đoàn Business",
        "Giám đốc phân phối cấp Quận",
        "Affiliate trực tiếp",
        "52,2 × 60% = 31,32 tỷ đ",
        "52,2 × 5% = 2,61 tỷ đ",
        "52,2 × 20% = 10,44 tỷ đ",
        "52,2 × 10% = 5,22 tỷ đ",
        "không chỉ là phúc lợi doanh nghiệp",
    ]
    for phrase in required:
        assert phrase in text, f"Missing required phrase: {phrase}"

    forbidden = [
        "Affiliate 2 tầng",
        "Affiliate hai tầng",
        "52,2 × 15% = 7,83 tỷ đ",
        "Cung ứng dịch vụ\n50%",
        "20% + 15% + 5% + 10% + 50%",
        "S12",
    ]
    for phrase in forbidden:
        assert phrase not in text, f"Stale phrase remains: {phrase}"

    with ZipFile(DOCX) as archive:
        assert archive.testzip() is None

    print("Allocation total: 100%")
    for role, value in expected.items():
        print(f"{role}: {value:,} VND")
    print(f"DOCX paragraphs: {len(document.paragraphs)}")
    print(f"DOCX tables: {len(document.tables)}")
    print(f"Inline images: {len(document.inline_shapes)}")
    print("Validation: READY TO SHARE")


if __name__ == "__main__":
    main()
